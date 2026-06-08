import json, os, openpyxl
M=json.load(open("/tmp/_model.json"))
AN="/home/user/puddlejumper/docs/analyses"; ROOT=f"{AN}/swanzey-operating-plan"; VIS=f"{ROOT}/Visuals"
CW=f"{AN}/swanzey-capital-forecast/Swanzey Capital Workbook.xlsx"; VER="v0.2"; TITLE="Swanzey Operating Plan"
# paths for projects
cw=openpyxl.load_workbook(CW); Pp=cw["Projects"]; paths={}
r=5
while Pp.cell(r,1).value and str(Pp.cell(r,1).value).upper()!="TOTAL":
    paths[Pp.cell(r,1).value]=Pp.cell(r,8).value; r+=1
def Mn(x): return f"${x/1e6:.2f}M"
def pct(x): return f"{x*100:.0f}%"
bl=M["backlog"]; lo=M["local"]; ntx=M["nontax"]
META=["PRIVATE working document · "+VER+" (auto-built from the workbooks) · June 2026",
      "Capital tables & charts generated live from the Capital Workbook; forecast from the Financial Model",
      "Edit the workbook input cells, re-run build, and this plan refreshes"]
SUB="Budget · Forecast · Capital — a consolidated operating reference"
C=[]
def h1(t):C.append(("h1",t))
def h2(t):C.append(("h2",t))
def p(t):C.append(("p",t))
def lead(a,b):C.append(("lead",a,b))
def bul(x):C.append(("bul",x))
def num(x):C.append(("num",x))
def table(h,r,w):C.append(("table",h,r,w))
def callout(t):C.append(("callout",t))
def img(f,c):C.append(("img",f,c))
def pb():C.append(("pb",))
C.append(("cover",TITLE,SUB,META))
h1("0.  Purpose")
p("One reference holding the Town's budget, a forecast, and the capital plan together — and it is now "
  "auto-built from the workbooks: the capital tables, the funding chart, the tax-glide chart, and the "
  "forecast are generated from the Capital Workbook and Financial Model. Edit the input cells, re-run "
  "the build, and this document refreshes. Every figure traces to the Town's own reports; capital "
  "figures and forecasts are clearly-labeled planning estimates.")
callout("Operating thesis: the Town will SAVE but won't BORROW. Fund the backlog with majority-vote "
        "reserves + state aid + the water revolving fund — not bonds — and the projects that keep "
        "stalling get built without a tax spike or a 3/5 vote.")
pb(); h1("PART I — BUDGET")
h2("1.  The 15-year record")
p("Operating budgets have risen ~55% since 2012 to roughly $8.94M (2026), with a flat early-2010s "
  "(~+1%/yr) giving way to a +6–7%/yr climb since 2020.")
img("01_budget-trajectory.png","Operating budget by town-meeting year, 2012–2026, with the forecast band (Part II).")
table(["Meeting year","Operating (proposed)","Default","Budget % YES"],
 [["2012","$5,747,885","$5,705,757","80%"],["2016","$6,262,426","$6,067,144","68%"],
  ["2020","$6,303,000","$6,072,735","—"],["2023","$7,425,000","$7,468,834","62%"],
  ["2025","$8,370,000","$7,935,699","57%"],["2026","$8,939,036","$8,891,785","passed"]],[1.6,2.0,1.7,1.4])
h2("2.  How the budget works (and why the 'no' lever is gone)")
bul([("Governing law — RSA 32.","Appropriation by purpose; no over-expenditure (32:8); transfers within the total (32:10); lapse (32:7); emergency over-spend needs DRA approval (32:11)."),
     ("Forms.","MS-636 / MS-737; MS-232 'Appropriations as Voted' filed within 20 days."),
     ("SB2 default — RSA 40:13, IX.","If the operating article fails, the default takes effect. In Swanzey it now sits within ~$47K of the proposal — and exceeded it in 2017, 2018, and 2023. 'Vote no to save' no longer works.")])
img("04_default-gap-collapse.png","Proposed minus default budget by meeting year — repeatedly near zero or negative.")
h2("3.  Cost drivers")
p("Structural, not programmatic: health insurance, wages, the NH Retirement System (RSA 100-A) employer "
  "rate, and the ambulance/EMS contract — largely non-discretionary.")
h2("4.  Tax rate & the 2024 revaluation")
p("Total rate: $24.51 (2021) → $26.87 (2023) → $19.06 (2024) → $20.18 (2025). The 2024 drop was a "
  "revaluation, not a tax cut. Education is ~60% of the bill; the municipal share is ~a quarter.")
img("05_tax-rate-components.png","Total tax rate by component, 2021–2025.")
pb(); h1("PART II — FORECAST")
p(f"Auto-built from the Financial Model (base ${M['forecast'][0]['base']:,.0f}; low/base/high growth). Scenarios, not predictions.")
frows=[]
for f in M["forecast"]:
    lbl=f"{f['yr']}"+(" (actual)" if f['yr']==2026 else "")
    frows.append([lbl,Mn(f['low']),Mn(f['base']),Mn(f['high'])])
table(["Fiscal year","Low (4.5%)","Base (6.0%)","High (7.5%)"],frows,[1.4,1.6,1.6,1.6])
p(f"Base case reaches {Mn(M['forecast'][-1]['base'])} by {M['forecast'][-1]['yr']} — purely on the cost escalators.")
h2("Budget-approval risk & fund balance")
p("Operating-budget support slid from ~80% (2011–15) to 57% (2025); a budget failure is plausible "
  "within a few cycles, though muted because the default ≈ the proposal.")
img("02_budget-support-erosion.png","Operating-budget support, 2011–2025, with trend.")
lead("Fund-balance guidance.","Hold unassigned fund balance in DRA's prudent band (≈5–17% of the "
     "operating budget); treat any amount applied to reduce taxes as multi-year smoothing.")
pb(); h1("PART III — CAPITAL PLAN")
h2("5.  The backlog and the 3/5 problem")
p(f"About {Mn(bl)} of identified needs. The public-works facility won majorities in 2024 and 2026 but "
  "failed the 3/5 a bond requires; reserves and credits pass 78–90%.")
img("03_reserves-vs-bonds.png","What passes easily versus what dies at the 3/5 wall.")
h2("6.  The funding strategy — other people's money")
p(f"About {pct(ntx/bl)} of the backlog is non-property-tax money (state bridge aid, DWSRF water rates, "
  f"grants), leaving {Mn(lo)} local — saved via majority-vote reserves, with $0 new bonds.")
img("06_capital-funding-stack.png","Funding stack by project (generated from the Capital Workbook).")
prows=[]
for pj in M["projects"]:
    prows.append([pj["name"],Mn(pj["total"]),Mn(pj["local"]),paths.get(pj["name"],"")])
prows.append(["TOTAL",Mn(bl),f"{Mn(lo)} ({pct(lo/bl)})","$0 new bonds"])
table(["Project","Total","Local share","Primary funding path"],prows,[1.9,1.0,1.3,2.4])
h2("7.  The reserve ladder & the tax glide")
lrows=[[x["fund"],f"${x['dep']:,.0f}",x["funds"],"majority"] for x in M["ladder"]]
table(["Capital Reserve Fund","Annual deposit","Funds","Vote"],lrows,[2.1,1.4,2.0,0.9])
img("07_capital-tax-glide.png",f"Net-new tax impact on a ${M['MED']/1000:.0f}K home: peaks ≈ ${round(M['peak_home'])}, then below today once built.")
h2("8.  Decision rules")
table(["If the project is…","Use…","Because"],
 [["≤ ~$250K, recurring","a capital reserve (pay-go)","simple majority, no interest"],
  ["a town bridge (≥10 ft)","RSA 234 aid + reserve for 20%","state pays 80%"],
  ["water/sewer","DWSRF, repaid by rates","off the property tax"],
  ["large, urgent, unavoidable","Bond Bank (RSA 33) — last resort","only when reserves can't get there in time"],
  ["anything","listed in the published CIP first","surprise becomes expectation"]],[1.7,2.4,2.4])
pb(); h1("PART IV — OPERATING PLAYBOOK")
h2("9.  The first twelve months")
table(["Quarter","Moves"],
 [["Q1","Stabilize the administrator's office (multi-year agreement + written evaluation). Publish a draft six-year CIP. Issue the one-page tax-bill explainer."],
  ["Q2","Warrant & pass the capital reserve ladder (majority votes). File State Bridge Aid for Christian Hill Rd. Commission the water rate study."],
  ["Q3","Advance DWSRF water financing. Deliver one visible small win (ADA via CDBG + reserve). Adopt operating norms for the five-member board."],
  ["Q4","Bring the full CIP + reserve plan to the deliberative session with an information night. Report the year's progress publicly."]],[0.9,5.7])
h2("10.  The annual municipal cycle")
table(["When","What"],
 [["By Sept 1","File MS-1 (valuation), MS-434 (revenues), MS-535 (financial report)"],
  ["Oct–Dec","Build budget & default; CIP update; budget public hearing; post warrant (RSA 31:95)"],
  ["~1st Tue Feb","SB2 deliberative session"],["~2nd Tue Mar","SB2 ballot: budget, articles, officers"],
  ["Meeting +20 days","File MS-232 'Appropriations as Voted'"],["Fall","DRA sets the tax rate; issue tax bills"]],[1.5,5.1])
h2("11.  Dashboard — numbers to watch")
table(["Metric","Now"],
 [["Capital backlog",f"{Mn(bl)} ({pct(ntx/bl)} non-tax; {Mn(lo)} local)"],
  ["Peak reserve tax impact",f"≈ ${round(M['peak_home'])}/yr on a ${M['MED']/1000:.0f}K home, then below today"],
  ["Operating budget / support","~$8.94M (2026); support 57% — arrest the slide with transparency"],
  ["Unassigned fund balance","keep ~5–17% of operating budget"],
  ["Total tax rate / valuation",f"$20.18 (2025); valuation input = ${M['VAL']/1e6:.1f}M (REFRESH with real MS-1)"],
  ["Next statutory deadline","always know it (MS-1 / MS-232 / hearing / warrant)"]],[2.0,4.6])
h1("Appendix — Provenance & sync")
p("Auto-built from: the Capital Workbook (Assumptions, Projects, Reserve Ladder, 10-Year Plan) and the "
  "Financial Model (Forecast). To refresh: edit the input cells in those workbooks, then run "
  "build_from_workbook.py + render_opplan.py. Real figures from the Town's annual reports (2011–2025), "
  "NH DRA, and local news; capital figures and forecasts are planning estimates. Key statutes: RSA 32, "
  "40:13, 35, 33/35-A, 234, 38, 162-K, 674:5-8, 91-A, 33-A, 165, 273-A, 100-A.")

# ===== renderer (shared) =====
from PIL import Image as PImage
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, ListFlowable, ListItem, HRFlowable, KeepTogether, Image)
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
NAVY=colors.HexColor("#1f3a5f"); GREY=colors.HexColor("#5f5e5a"); LT=colors.HexColor("#eef1f4"); GOLD=colors.HexColor("#b8860b")
PH1=ParagraphStyle("PH1",fontSize=15,textColor=NAVY,fontName="Helvetica-Bold",spaceBefore=4,spaceAfter=6)
PH2=ParagraphStyle("PH2",fontSize=11.5,textColor=NAVY,fontName="Helvetica-Bold",spaceBefore=8,spaceAfter=3)
PBD=ParagraphStyle("PBD",fontSize=10.2,leading=14.3,spaceAfter=6,alignment=TA_JUSTIFY)
PBUL=ParagraphStyle("PBUL",fontSize=9.9,leading=13.2,spaceAfter=3)
def asp(fp): iw,ih=PImage.open(fp).size; return ih/iw
def render_pdf():
    f=[]
    for blk in C:
        k=blk[0]
        if k=="cover":
            _,t,s,meta=blk
            f+=[Spacer(1,1.7*inch),Paragraph(t,ParagraphStyle("ct",fontSize=30,textColor=NAVY,fontName="Helvetica-Bold",alignment=TA_CENTER,leading=34)),Spacer(1,8),HRFlowable(width="55%",thickness=1.3,color=GOLD,hAlign="CENTER"),Spacer(1,12),Paragraph(s,ParagraphStyle("cs",fontSize=13,textColor=GREY,alignment=TA_CENTER,leading=18)),Spacer(1,1.3*inch)]
            for m in meta: f.append(Paragraph(m,ParagraphStyle("cm",fontSize=10,textColor=GREY,alignment=TA_CENTER,spaceAfter=3)))
            f.append(PageBreak())
        elif k=="h1": f.append(Paragraph(blk[1],PH1))
        elif k=="h2": f.append(Paragraph(blk[1],PH2))
        elif k=="p": f.append(Paragraph(blk[1],PBD))
        elif k=="lead": f.append(Paragraph(f"<b>{blk[1]}</b> {blk[2]}",PBD))
        elif k=="bul":
            f+=[ListFlowable([ListItem(Paragraph((f"<b>{x[0]}</b> {x[1]}" if isinstance(x,tuple) else x),PBUL),value="•",leftIndent=12) for x in blk[1]],bulletType="bullet"),Spacer(1,5)]
        elif k=="img":
            fp=f"{VIS}/{blk[1]}"; w=5.6
            f.append(KeepTogether([Image(fp,width=w*inch,height=w*asp(fp)*inch),Paragraph(blk[2],ParagraphStyle("cap",fontSize=7.7,textColor=GREY,alignment=TA_CENTER,spaceBefore=2,spaceAfter=9,fontName="Helvetica-Oblique"))]))
        elif k=="table":
            hd,rows,ws=blk[1],blk[2],blk[3]
            data=[[Paragraph(f"<b>{c}</b>",ParagraphStyle('th',fontSize=8.6,textColor=colors.white)) for c in hd]]+[[Paragraph(str(c),ParagraphStyle('td',fontSize=8.7,leading=11)) for c in r] for r in rows]
            t=Table(data,colWidths=[x*inch for x in ws],repeatRows=1)
            t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),NAVY),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,LT]),("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#c8ccd2")),("VALIGN",(0,0),(-1,-1),"TOP"),("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),("LEFTPADDING",(0,0),(-1,-1),4)]))
            f+=[t,Spacer(1,8)]
        elif k=="callout":
            t=Table([[Paragraph(blk[1],ParagraphStyle('co',fontSize=9.9,leading=13.8,textColor=NAVY))]],colWidths=[6.5*inch])
            t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),LT),("BOX",(0,0),(-1,-1),0.6,NAVY),("LINEBEFORE",(0,0),(0,-1),3,GOLD),("TOPPADDING",(0,0),(-1,-1),7),("BOTTOMPADDING",(0,0),(-1,-1),7),("LEFTPADDING",(0,0),(-1,-1),9),("RIGHTPADDING",(0,0),(-1,-1),9)]))
            f+=[Spacer(1,2),t,Spacer(1,8)]
        elif k=="pb": f.append(PageBreak())
    return f
def foot(c,d):
    c.saveState(); c.setFont("Helvetica",7.5); c.setFillColor(GREY)
    c.drawString(0.85*inch,0.5*inch,f"Swanzey Operating Plan · {VER} · PRIVATE · auto-built")
    c.drawRightString(7.65*inch,0.5*inch,str(d.page)); c.restoreState()
SimpleDocTemplate(f"{ROOT}/Swanzey Operating Plan {VER}.pdf",pagesize=letter,topMargin=0.7*inch,bottomMargin=0.75*inch,leftMargin=0.85*inch,rightMargin=0.85*inch,title=TITLE).build(render_pdf(),onLaterPages=foot,onFirstPage=lambda c,d:None)
print("PDF",VER,"done")
# DOCX
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
dN=RGBColor(0x1f,0x3a,0x5f); dG=RGBColor(0x5f,0x5e,0x5a)
doc=Document()
for s in doc.sections: s.top_margin=Inches(0.7); s.bottom_margin=Inches(0.7); s.left_margin=Inches(0.9); s.right_margin=Inches(0.9)
doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(10.5)
def dp(a=6,al=None):
    pp=doc.add_paragraph(); pp.paragraph_format.space_after=Pt(a)
    if al is not None: pp.alignment=al
    return pp
def dr(pp,t,sz=10.5,b=False,col=None,it=False):
    r=pp.add_run(t); r.bold=b; r.italic=it; r.font.size=Pt(sz)
    if col is not None: r.font.color.rgb=col
for blk in C:
    k=blk[0]
    if k=="cover":
        _,t,s,meta=blk
        for _ in range(3): doc.add_paragraph()
        dr(dp(6,WD_ALIGN_PARAGRAPH.CENTER),t,28,True,dN); dr(dp(16,WD_ALIGN_PARAGRAPH.CENTER),s,13,col=dG)
        for _ in range(2): doc.add_paragraph()
        for m in meta: dr(dp(2,WD_ALIGN_PARAGRAPH.CENTER),m,10,col=dG)
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    elif k=="h1": dr(dp(4),blk[1],15,True,dN)
    elif k=="h2": dr(dp(2),blk[1],12,True,dN)
    elif k=="p": pp=dp(); dr(pp,blk[1]); pp.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    elif k=="lead": pp=dp(); dr(pp,blk[1]+" ",10.3,True); dr(pp,blk[2],10.3); pp.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    elif k=="bul":
        for x in blk[1]:
            pp=doc.add_paragraph(style="List Bullet"); pp.paragraph_format.space_after=Pt(3)
            if isinstance(x,tuple): dr(pp,x[0]+" ",10.2,True); dr(pp,x[1],10.2)
            else: dr(pp,x,10.2)
    elif k=="img":
        doc.add_picture(f"{VIS}/{blk[1]}",width=Inches(5.9)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        dr(dp(8,WD_ALIGN_PARAGRAPH.CENTER),blk[2],7.7,it=True,col=dG)
    elif k=="table":
        hd,rows,ws=blk[1],blk[2],blk[3]
        t=doc.add_table(rows=1,cols=len(hd)); t.style="Light Grid Accent 1"
        for i,c in enumerate(hd):
            rr=t.rows[0].cells[i].paragraphs[0].add_run(c); rr.bold=True; rr.font.size=Pt(9)
        for row in rows:
            cells=t.add_row().cells
            for i,c in enumerate(row):
                rr=cells[i].paragraphs[0].add_run(str(c)); rr.font.size=Pt(9)
        dp(6)
    elif k=="callout":
        t=doc.add_table(rows=1,cols=1); t.style="Light Shading Accent 1"
        rr=t.rows[0].cells[0].paragraphs[0].add_run(blk[1]); rr.font.size=Pt(10); rr.font.color.rgb=dN
        dp(6)
    elif k=="pb": doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
doc.save(f"{ROOT}/Swanzey Operating Plan {VER}.docx"); print("DOCX",VER,"done")
