#!/usr/bin/env python3
"""
PublicLogic — Five Tier-1 outreach emails. Short, plain, no hype, no pitch deck.
Each offers a paid Stewardship Map on a specific stalled function. Grounded in
real work. Placeholders in [brackets] to personalize. Copy/paste-ready .docx.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SLATE=RGBColor(0x1A,0x1D,0x20); GREEN=RGBColor(0x0F,0x4C,0x3A)
GOLD=RGBColor(0xA9,0x77,0x2F); MUTE=RGBColor(0x6B,0x66,0x60)
SERIF="Georgia"; SANS="Calibri"

d=Document()
for s in d.sections:
    s.top_margin=Inches(0.6); s.bottom_margin=Inches(0.55)
    s.left_margin=Inches(0.9); s.right_margin=Inches(0.9)

def para(text,font=SANS,size=10.5,color=SLATE,bold=False,ital=False,before=2,after=2,sp=1.08):
    p=d.add_paragraph()
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    r=p.add_run(text); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=color; f.bold=bold; f.italic=ital
    return p

def rule():
    p=d.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(4)
    r=p.add_run("—"*60); r.font.color.rgb=RGBColor(0xCC,0xC8,0xC0); r.font.size=Pt(9)

def email(n,who,when,subject,body,note):
    para(f"EMAIL {n} · {who}",SANS,11,GREEN,bold=True,before=8,after=0)
    para(when,SANS,9,GOLD,bold=True,ital=True,before=0,after=2)
    para("Subject:  "+subject,SANS,10.5,SLATE,bold=True,after=3)
    for line in body:
        if line=="":
            para("",SANS,5,SLATE,before=0,after=0)
        else:
            para(line,SANS,10.5,SLATE,after=0)
    para("Why this one works:  "+note,SANS,9,MUTE,ital=True,before=4,after=2)
    rule()

# ---- Header
para("PUBLICLOGIC",SANS,11,GOLD,bold=True,after=0)
para("Five Tier-1 Outreach Emails",SERIF,22,SLATE,bold=True,before=0,after=0)
para("Short and plain on purpose. No deck, no pitch. Each one asks for a small, specific first step: a paid "
     "Stewardship Map on one function that's already stuck. Personalize the [brackets] before sending.",
     SANS,10,GREEN,ital=True,before=1,after=2)
para("Rule of thumb: keep them this short. The goal of the email is a 20-minute call, not a sale.",
     SANS,9.5,MUTE,ital=True,after=2)
rule()

email(1,"Warm municipal contact (Sutton / Shrewsbury / Phillipston)","Relationship: Strongest · pursue first",
 "A small idea for [Town] — quick call?",
 ["Hi [Name],",
  "",
  "Good to be back in touch. I've started a practice, PublicLogic, focused on a problem we both know well: "
  "the work that lives in one person's head and quietly breaks when they leave or get pulled onto something else.",
  "",
  "We do a short, fixed-fee piece called a Stewardship Map — a 2–4 week look at how one function actually "
  "runs (records, who owns what, where it would stall) and the next moves to shore it up. It's small on "
  "purpose; it's the safe way to start.",
  "",
  "Is there one function in [Town] you've worried about — [grants / capital planning / a department that "
  "runs through one person] — that would be worth mapping? Happy to grab 20 minutes and tell you if it's "
  "even a fit.",
  "",
  "Best,",
  "Nathan"],
 "Strongest relationship, so it's the most direct. Names a real pain and a small, bounded ask. No jargon."),

email(2,"CMRPC — Central Mass Regional Planning Commission","Relationship: Developing · warm intro ideal",
 "Stewardship Maps for your member towns?",
 ["Hi [Name],",
  "",
  "I run PublicLogic — I work with municipalities on continuity: making sure the records, responsibilities, "
  "and know-how behind a function survive turnover. A lot of it overlaps with the capacity gaps CMRPC sees "
  "in member towns every day.",
  "",
  "We offer a short, fixed-fee Stewardship Map ($2,500–$7,500) that makes one function visible and hands the "
  "town a clear next step. It can stand alone or sit alongside the work CMRPC already does.",
  "",
  "Would it be worth a short call to see whether there's a town or two where this would help? I'm glad to "
  "walk through a quick example.",
  "",
  "Thanks,",
  "Nathan"],
 "Positions PublicLogic as complementary to the RPC, not competitive. Offers value to THEIR members."),

email(3,"MRPC — Montachusett Regional Planning Commission","Relationship: Developing",
 "A continuity tool that might help a few member towns",
 ["Hi [Name],",
  "",
  "I'm Nathan Boudreau — municipal background (MPA/MCPPO), now running PublicLogic. We help towns keep the "
  "knowledge and records behind a function from walking out the door when someone leaves.",
  "",
  "Our entry piece is a Stewardship Map: a short, fixed-fee diagnostic of one function — how it runs, where "
  "it's fragile, what to do next. Towns like it because it's small, concrete, and low-risk.",
  "",
  "If a couple of MRPC towns come to mind, I'd value 20 minutes to see if it's a fit. No pitch — just a quick "
  "look at where it helps.",
  "",
  "Best,",
  "Nathan"],
 "Leads with credibility (MPA/MCPPO), then the same small ask. Mirrors the CMRPC note without copying it."),

email(4,"Existing network (general warm contact)","Relationship: Strongest · ask directly",
 "Who do you know who needs this?",
 ["Hi [Name],",
  "",
  "Quick one. I've launched PublicLogic — we help organizations keep important work from breaking when the "
  "person who held it moves on. The entry point is a short, fixed-fee Stewardship Map of one function.",
  "",
  "You know my work. Who's the first person you'd point me to — someone with a function that's stuck, "
  "fragile, or running entirely through one person?",
  "",
  "An intro would mean a lot. Happy to send a one-pager you can forward.",
  "",
  "Thanks,",
  "Nathan"],
 "Doesn't sell — asks for a referral. The warmest contacts open doors faster than they buy."),

email(5,"Engineering / municipal consulting firm (Fuss & O'Neill, Community Paradigm)","Relationship: Cold · partnership angle",
 "A partner for the continuity side of your municipal work",
 ["Hi [Name],",
  "",
  "I run PublicLogic. We handle a narrow but recurring piece of municipal work: the continuity layer — "
  "records, ownership, and institutional knowledge behind a function — so projects don't stall when staff "
  "turn over.",
  "",
  "It tends to sit right next to what your firm already delivers. Where you build the plan or the system, "
  "we make sure the town can actually carry it after handoff. Our Stewardship Map is a clean, fixed-fee way "
  "to start on a shared client.",
  "",
  "Would a short call make sense to see where our work could complement yours?",
  "",
  "Best,",
  "Nathan"],
 "Cold, so it's framed as partnership, not competition — 'next to,' 'complement,' 'after handoff.'"),

para("After any reply:  offer the relevant Scope & Fee sheet (Stewardship Map or Permit Path Scan) and propose "
     "a specific function to map. Keep the first call to listening — name the stuck function in their words.",
     SANS,9.5,GREEN,ital=True,before=6,after=2)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"final_deck","PublicLogic - Five Outreach Emails (Tier 1).docx")
os.makedirs(os.path.dirname(out),exist_ok=True)
d.save(out); print("wrote",out)
