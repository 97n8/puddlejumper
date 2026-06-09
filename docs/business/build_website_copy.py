#!/usr/bin/env python3
"""
PublicLogic — Complete v1 Website Copy (whole ecosystem).
Every page, ready to hand to a web builder: slug, title, meta, hero, body,
CTAs. Canon voice — grounded, no hype — and obeys the Do-Not-Say guardrails.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SLATE=RGBColor(0x1A,0x1D,0x20); GREEN=RGBColor(0x0F,0x4C,0x3A)
GOLD=RGBColor(0xA9,0x77,0x2F); MUTE=RGBColor(0x6B,0x66,0x60)
SERIF="Georgia"; SANS="Calibri"
d=Document()
for s in d.sections:
    s.top_margin=Inches(0.65); s.bottom_margin=Inches(0.6); s.left_margin=Inches(0.9); s.right_margin=Inches(0.9)

def P(text,font=SANS,size=10.5,color=SLATE,bold=False,ital=False,before=2,after=2,sp=1.08,align=WD_ALIGN_PARAGRAPH.LEFT):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    r=p.add_run(text); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=color; f.bold=bold; f.italic=ital
    return p
def page(title,slug,meta):
    d.add_page_break()
    P(title,SERIF,20,SLATE,bold=True,before=0,after=0)
    P("URL: "+slug+"     ·     Meta: "+meta,SANS,8.5,MUTE,ital=True,before=1,after=4)
def lab(text): P(text,SANS,9.5,GOLD,bold=True,before=8,after=1)   # field label (HERO, BODY, CTA)
def h(text): P(text,SANS,12,GREEN,bold=True,before=7,after=1)     # on-page section heading
def body(text): P(text,SANS,10.5,SLATE,after=2)
def hero(text): P(text,SERIF,16,SLATE,bold=True,after=1,sp=1.05)
def sub(text): P(text,SANS,11.5,GREEN,ital=True,after=2)
def bullets(items,g=GOLD):
    for it in items:
        p=d.add_paragraph(); p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1.06
        r=p.add_run("—  "); r.font.name=SANS; r.font.size=Pt(10.5); r.font.color.rgb=g; r.font.bold=True
        r2=p.add_run(it); r2.font.name=SANS; r2.font.size=Pt(10.5); r2.font.color.rgb=SLATE
def cta(text): P("CTA button:  "+text,SANS,10.5,SLATE,bold=True,before=3,after=2)

# ============ COVER ============
P("PUBLICLOGIC",SANS,11,GOLD,bold=True,after=0)
P("Website Copy — v1 Ecosystem",SERIF,24,SLATE,bold=True,before=0,after=0)
P("Every page, ready for a web builder. Plain, grounded voice — outcomes, never hype. "
  "Obeys the language guardrails (no “AI,” no “software platform,” no “dashboard as the product”). "
  "June 2026 · Private & Confidential.",SANS,10,GREEN,ital=True,before=2,after=6)
h("Sitemap")
bullets(["Home  /","How It Works  /how-it-works","LogicCommons  /logiccommons","Permit & Bridge  /permit-and-bridge",
 "Stewardship Map  /stewardship-map","Services  /services","Proof  /proof","About  /about",
 "FAQ  /faq","Start  /start  (contact)","Global elements (nav · footer · CTAs)"])
P("Voice rules carried throughout: lead with the problem, tie every mention of stewardship to a real pain, "
  "name outcomes not protocols, and keep the public offer as Continuity & Stewardship Systems — the plumbing, "
  "not a product.",SANS,9.5,MUTE,ital=True,before=6,after=2)

# ============ GLOBAL ============
page("Global Elements","(nav / footer / reusable)","Applies to every page")
lab("WORDMARK"); body("PublicLogic")
lab("TAGLINE (beneath wordmark where helpful)"); body("Continuity • Data • Stewardship")
lab("PRIMARY NAV"); body("Home   ·   How It Works   ·   LogicCommons   ·   Services   ·   Proof   ·   About   ·   Start")
lab("GLOBAL CTA (header button)"); body("Start with a Stewardship Map")
lab("FOOTER — TAGLINE LINE"); P("Honor the past. Improve the present. Continue the work.",SERIF,12,GREEN,bold=True,ital=True,after=2)
lab("FOOTER — BOILERPLATE")
body("PublicLogic helps organizations understand, improve, and continue their work. Municipal government is "
     "our home ground, but the method serves mission-driven and legacy-minded organizations wherever continuity matters.")
lab("FOOTER — CONTACT"); body("[hello@publiclogic.org]   ·   [978-807-0829]   ·   PublicLogic LLC")
lab("FOOTER — FINE PRINT")
body("PublicLogic provides stewardship and advisory services. We are not a software vendor and our materials "
     "are not legal advice. © PublicLogic LLC. All rights reserved.")

# ============ HOME ============
page("Home","/","Institutional stewardship that keeps important work alive through turnover and change. Start with a Stewardship Map.")
lab("HERO HEADLINE"); hero("Most organizations don’t lose ground from a lack of effort. They lose it from a lack of visibility.")
lab("HERO SUBHEAD"); sub("PublicLogic helps projects move through public systems — and helps organizations keep the capacity to sustain what they build.")
lab("HERO CTA"); cta("Start with a Stewardship Map"); body("Secondary link: See how it works →")
lab("ONE-LINE VERSION (band under hero)")
P("Help people understand the path. Help projects move through the path. Leave the path easier for the next person.",
  SERIF,13,GREEN,bold=True,ital=True,after=2)

h("The quiet problem")
body("Good work fails for boring reasons. The person who held it leaves. How and why things were done was never "
     "written down. A funding window closes before anyone is ready. A plan gets made, then sits on a shelf. By the "
     "time anyone notices, the momentum, the money, or the knowledge is already gone.")
body("The central question isn’t whether risk exists. It’s where it will live — and who carries it when the people change.")

h("How we help — five steps, in plain language")
bullets([
 "Understand — free public tools that explain the path (LogicCommons).",
 "Navigate — help finding the way through permits, boards, and funding (Permit & Bridge).",
 "Diagnose — a short, paid look at what actually needs to happen (Stewardship Map).",
 "Deliver — the project, funding, or capacity work itself (PublicLogic).",
 "Sustain — the records and routines that keep it working after we step back (Continuity & Stewardship Systems).",
])
body("You can enter anywhere. Most clients start with a Stewardship Map — the safe, fixed-fee first step.")

h("What makes us different")
body("Most systems manage work. PublicLogic helps steward what has to survive the work. And we have a rule we hold "
     "ourselves to: we do not create dependency. We create understanding.")

h("Who it’s for")
body("Municipal governments and the partners around them — regional planning agencies, engineering firms, and "
     "municipal consulting firms — plus mission-driven and legacy-minded organizations wherever continuity matters.")

lab("CLOSING CTA SECTION")
P("If your planner leaves, does this still work?",SERIF,15,SLATE,bold=True,after=1)
body("That’s the question a Stewardship Map answers. Small, fixed-fee, and honest.")
cta("Start with a Stewardship Map")

# ============ HOW IT WORKS ============
page("How It Works","/how-it-works","The PublicLogic path: Understand, Navigate, Diagnose, Deliver, Sustain — and the Continuity & Stewardship Systems underneath.")
lab("HERO"); hero("One path, five steps. Enter anywhere.")
lab("SUBHEAD"); sub("You don’t have to buy the whole thing to start. Each step stands on its own and leads naturally to the next.")

h("The path")
bullets([
 "Understand · LogicCommons — free templates, checklists, and frameworks that help you see the path before you need paid help.",
 "Navigate · Permit & Bridge — help moving through permits, boards, funding, and stakeholders.",
 "Diagnose · Stewardship Map — a short, fixed-fee look at how one function really runs and what to do next.",
 "Deliver · PublicLogic — the project development, funding, implementation, or capacity work itself.",
 "Sustain · Continuity & Stewardship Systems — the records and routines that keep it all working after we step back.",
])

h("How we actually do the work")
body("Our method is four moves: Map → Embed → Encode → Sustain. We map how work, knowledge, authority, records, "
     "and risk actually move; we embed the right structure into the live workflow with the people who do the work; "
     "we encode the knowledge and rules so they outlast any one person; and we sustain it until it runs without us.")
body("Across an engagement, a client moves through a simple cycle: Discover · Honor · Understand · Improve · Build · "
     "Steward · Continue. We honor what already works before we change anything.")

h("What “Continuity & Stewardship Systems” means")
body("It’s the plumbing, not a product. Continuity & Stewardship Systems are the records, processes, templates, "
     "environments, accountability structures, and proof that help important work survive turnover, changing "
     "priorities, and organizational change. Continuity is the outcome. Stewardship is the practice. Systems are the "
     "mechanism. You never have to think about the parts — you just get work that holds together.")
cta("Start with a Stewardship Map")

# ============ LOGICCOMMONS ============
page("LogicCommons","/logiccommons","Free, plain-language tools that help residents, owners, and organizations understand the path before they need paid help.")
lab("HERO"); hero("Help understanding the path — free.")
lab("SUBHEAD"); sub("LogicCommons is our public layer: free templates, checklists, and frameworks. No account, no pitch. Just help.")
h("Why it’s free")
body("Some questions don’t need a consultant. They need a checklist, a starting point, and the right office to call. "
     "LogicCommons exists for that. Its purpose is access, trust, and triage — not selling. It does not replace your "
     "municipality, inspector, planner, or permitting authority.")
h("Start here")
bullets([
 "“Can I Do This?” — a plain-language permit-path worksheet for residents, small owners, and project sponsors.",
 "More templates and checklists added over time, drawn from real municipal work.",
])
cta("Download the “Can I Do This?” worksheet")
h("When you need more than a template")
body("If your project touches several systems at once, involves a public hearing, or you’ve already been bounced "
     "between offices, a Permit Path Scan ($250–$750) lays out the likely path, the boards involved, and the "
     "documents you’ll need — so you stop guessing.")
P("LogicCommons helps people start. PublicLogic helps them carry it through.",SERIF,12,GREEN,bold=True,ital=True,before=4,after=2)
cta("See Permit & Bridge")

# ============ PERMIT & BRIDGE ============
page("Permit & Bridge","/permit-and-bridge","Help understanding the path, and help moving through it — for residents and for project sponsors.")
lab("HERO"); hero("Permit & Bridge helps the public understand the path and helps project sponsors move through it.")
lab("SUBHEAD"); sub("Two sides of the same bridge: simple help for the public, white-glove help for project sponsors.")
h("Two sides of the bridge")
bullets([
 "Public Help Layer — for residents, small owners, and nonprofits. “Can I do X in my backyard, on my property, in my building?” Cheap, simple, framework-guided.",
 "Project Sponsor Layer — for developers, towns, nonprofits, and businesses. A white-glove path through permits, boards, funding, and stakeholders.",
])
h("Why it exists")
body("Most sponsors understand their project. What stalls them is the public system around it — the permits, the "
     "boards, the funding rules, the politics. That gap is where good projects die. We become the bridge: we "
     "translate between the project and the public system so the project keeps moving and nothing falls through the "
     "cracks between the people involved.")
h("The offer ladder")
bullets([
 "Tier 0 · Public Permit Helper — free or very low cost.",
 "Tier 1 · Permit Path Scan — $250–$750.",
 "Tier 2 · Stewardship Map — $2,500–$7,500.",
 "Tier 3 · Permit & Bridge Sprint — $7,500–$15,000.",
 "Tier 4 · White-Glove Implementation — $3,500–$8,500 / month.",
 "Tier 5 · Funding / Grant Build — $5,000–$25,000, scope-dependent.",
])
body("The free tier builds trust; the paid tiers handle complexity. A paid tier credits toward the next "
     "engagement — credits don’t stack, so you never pay twice for the same step.")
h("Our promise to a sponsor")
P("You leave with a clear path forward, a clear record of how you got there, and a clear understanding of who owns "
  "what next.",SERIF,12,SLATE,bold=True,ital=True,after=2)
cta("Book a Permit & Bridge conversation")

# ============ STEWARDSHIP MAP ============
page("Stewardship Map","/stewardship-map","The safe first step: a short, fixed-fee look at how one function really runs, where it’s fragile, and what to do next.")
lab("HERO"); hero("The safest way to start working together.")
lab("SUBHEAD"); sub("A short, fixed-fee diagnostic that makes one function visible before anyone commits real money.")
h("What it is")
body("Before you commit to a project, a build, or a hire, the Stewardship Map shows how work, knowledge, authority, "
     "records, and risk actually move through your organization — and where the next person would get stuck.")
h("What you get")
bullets([
 "Stewardship Map — a clear picture of how the function runs today.",
 "Continuity Risks — the points where turnover or a gap would actually break something.",
 "Readiness Findings — an honest read on what your organization can absorb and sustain now.",
 "Priority Roadmap — a short, sequenced list of next moves, with what each is worth.",
 "Recommended Next Step — our honest read on the single best move, or whether there isn’t one yet.",
])
h("Fee")
body("$2,500–$7,500, fixed and non-contingent — our fee never depends on a funding outcome or a decision going a "
     "particular way. Typically 2–4 weeks. The Map credits toward a follow-on engagement if you continue.")
h("Why start here")
body("Low risk. Small, fixed, and bounded — you learn what’s really going on before committing to anything bigger. "
     "Most clients use the Map to decide what to do next.")
lab("PROOF LINE"); P("Similar work has been used to identify continuity, governance, and implementation risks before "
     "capital, technology, and organizational investments.",SANS,10,SLATE,ital=True,after=2)
cta("Start with a Stewardship Map")

# ============ SERVICES ============
page("Services","/services","From a free worksheet to ongoing capacity support — a clear ladder, honestly priced, non-contingent.")
lab("HERO"); hero("Simple offers, honestly priced.")
lab("SUBHEAD"); sub("Everything starts with making the work visible. The Map comes first; the rest follows from what it finds.")
h("The offers")
bullets([
 "Stewardship Map (entry) — make one function visible. $2,500–$7,500.",
 "Project Development Sprint — turn a priority or challenge into an actionable project. $5,000–$15,000.",
 "Funding Strategy Sprint — find realistic funding and the readiness it requires. $5,000–$12,500.",
 "Permit & Bridge Sprint — move a project through the public systems around it. $7,500–$15,000.",
 "Diagnostic — a deep review of risks, records, structure, and next steps when a Map reveals serious drift. $10,000–$25,000+.",
 "Implementation Support — coordinate and hold accountability while your team executes. $3,500–$8,500 / month.",
 "Capacity Support — step in and hold a role directly until it transfers back. $4,000–$10,000 / month.",
])
h("How we price")
body("Fixed-fee or retainer, and non-contingent. A paid tier credits toward the next engagement — credits don’t "
     "stack, so you never pay twice for the same step. We tell you the smallest useful next step, not the biggest.")
h("Implementation vs. Capacity")
body("Implementation Support helps your team execute — the work stays theirs. Capacity Support means we hold the "
     "role ourselves until it transfers back. Same discipline; different level of who holds the chair.")
cta("Start with a Stewardship Map")

# ============ PROOF ============
page("Proof","/proof","Stewardship is provable, not just stated: continuity, record, readiness, adoption, outcome.")
lab("HERO"); hero("We can show it.")
lab("SUBHEAD"); sub("Five kinds of proof, each tied to work we’ve actually done.")
h("How we prove the work")
bullets([
 "Continuity — a function survives when the person who ran it leaves. (Successor tests and structure mapped to authority.)",
 "Record — every claim, number, and decision traces to a source. (A governed record: source file → workbook row → output.)",
 "Readiness — the organization can absorb and sustain the change, assessed before any build.",
 "Adoption — staff actually use the system, monitored through implementation.",
 "Outcome — the work moved forward and got funded, documented.",
])
h("Where it comes from")
body("PublicLogic came out of real work — municipal administration, capital planning, grant development, and "
     "engagements across Massachusetts and beyond. We share specific examples in conversation, with client "
     "permission, rather than publishing confidential detail here.")
P("Note: some of our work is client-confidential. We confirm what we’re comfortable naming before sharing it.",
  SANS,9,MUTE,ital=True,after=2)
cta("Talk to us about your situation")

# ============ ABOUT ============
page("About","/about","PublicLogic emerged from the work — municipal administration, capital planning, grant development — not from a theory.")
lab("HERO"); hero("We didn’t start from a theory. We started from the work.")
lab("SUBHEAD"); sub("PublicLogic is the company that already existed in our day jobs — named, and made repeatable.")
h("Why we exist")
body("Across years inside municipal administration, capital planning, and grant development, we kept seeing the same "
     "thing: good, important work would stall or quietly disappear — not because the idea was wrong, but because the "
     "person holding it left, the knowledge was never written down, or no one owned what came next. PublicLogic is "
     "our attempt to fix that, with the practices we wish we’d had when we were inside it.")
h("Nathan + Allie")
body("Nathan Boudreau (MPA, MCPPO) knows institutional systems — how municipalities, money, grants, procurement, "
     "and projects actually work from the inside. Allie Weiss Rothschild knows human systems — how people, roles, "
     "adoption, and change actually work, and why systems get used or ignored. In our model, that’s the Governance "
     "Steward lane: the human-centered guide who makes systems fit real people and keeps them trusted and adopted "
     "over time. Most firms have one side or the other. Projects fail for whichever side gets ignored.")
h("What we believe")
bullets([
 "Vision: organizations should be easier to inherit than they were to build.",
 "Promise: we do not create dependency. We create understanding.",
 "North Star: every system should make it easier for the next person to do the right thing.",
])
cta("Start with a Stewardship Map")

# ============ FAQ ============
page("FAQ","/faq","Plain answers: what a Stewardship Map is, what it costs, and what PublicLogic is and isn’t.")
def qa(q,a):
    P("Q.  "+q,SANS,11,GREEN,bold=True,before=6,after=1); body(a)
lab("HERO"); hero("Straight answers.")
qa("Is this software?","No. PublicLogic is a stewardship and advisory practice. We work in ordinary tools your team "
   "already uses — documents, workbooks, registers, calendars — and make them behave like a governed system. The "
   "value is the discipline, not a platform.")
qa("Is this an “AI” service?","No. We use good judgment, real experience, and practical tools. We’re careful not to "
   "overclaim — we make the work easier to understand and continue, not magically automated.")
qa("What’s a Stewardship Map?","A short, fixed-fee diagnostic of one function: how it runs, where it’s fragile, and "
   "what to do next. It’s the safest way to start, because it tells both of us exactly what should come next.")
qa("What does it cost?","The Stewardship Map is $2,500–$7,500, fixed and non-contingent. Other offers range from a "
   "$250–$750 Permit Path Scan up to monthly support. Our fee never depends on a funding outcome.")
qa("How long does it take?","A Stewardship Map is usually 2–4 weeks. A Permit Path Scan is a few business days.")
qa("Do you replace our staff?","Only if you ask us to, and only until it transfers back. Capacity Support means we "
   "hold a role temporarily; Implementation Support means we help your team carry it. Either way, the goal is to "
   "make you less dependent on us, not more.")
qa("Who do you work with?","Municipal governments and their partners — regional planning agencies, engineering "
   "firms, municipal consulting firms — and mission-driven or legacy-minded organizations where continuity matters.")
qa("Where do I start?","With a Stewardship Map, or a short call. If you just have a permitting question, start free "
   "in LogicCommons.")
cta("Start with a Stewardship Map")

# ============ START / CONTACT ============
page("Start","/start","Start with a Stewardship Map, or just ask a question. A short, honest first conversation.")
lab("HERO"); hero("Start with a Stewardship Map.")
lab("SUBHEAD"); sub("Or just tell us what’s stuck. The first conversation is listening, not selling.")
h("Tell us a little")
lab("FORM FIELDS")
bullets([
 "Name",
 "Organization",
 "Email",
 "What’s the function or project you’re worried about? (a few sentences)",
 "Where are you — town, region, or organization type?",
 "What would “this is handled” look like to you?",
])
lab("FORM BUTTON"); cta("Send")
h("Or reach us directly")
body("[hello@publiclogic.org]   ·   [978-807-0829]")
P("If your planner leaves, does this still work? Let’s find out together.",SERIF,13,GREEN,bold=True,ital=True,before=4,after=2)

# ============ CTA LIBRARY ============
page("CTA & Microcopy Library","(reusable)","Reusable buttons and short lines")
lab("PRIMARY CTAs")
bullets(["Start with a Stewardship Map","Book a Permit & Bridge conversation","Download the “Can I Do This?” worksheet",
 "See how it works","Talk to us about your situation"])
lab("MICROCOPY LINES (use sparingly, always tied to the problem)")
bullets([
 "If your planner leaves, does this still work?",
 "Help the work survive the people who do it.",
 "We do not create dependency. We create understanding.",
 "Most systems manage work. PublicLogic helps steward what has to survive the work.",
 "Honor the past. Improve the present. Continue the work.",
])
lab("DO NOT USE (guardrails)")
bullets([
 "“AI-powered / AI transformation” → say Continuity & Stewardship Systems / institutional stewardship.",
 "“Software platform / SaaS” → say hosted operating environment.",
 "“Revolutionary / disruptive” → say serious, useful, durable.",
 "“The dashboard” as the product → say the Stewardship Map / the operating environment.",
 "“Institutional Stewardship” by itself → tie it to preventing turnover, lost knowledge, stalled projects, and missed opportunities.",
])

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"final_deck","PublicLogic - Website Copy (v1 Ecosystem).docx")
os.makedirs(os.path.dirname(out),exist_ok=True)
d.save(out); print("wrote",out)
