#!/usr/bin/env python3
"""
PublicLogic — Proof & Case Example sheet. Closes the "proof is thin" gap
honestly: a proof framework with the metric to capture for each, plus a
reusable one-page Case Example template to fill in after each engagement.
No invented numbers — placeholders the work will replace.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SLATE=RGBColor(0x1A,0x1D,0x20); GREEN=RGBColor(0x0F,0x4C,0x3A)
GOLD=RGBColor(0xA9,0x77,0x2F); MUTE=RGBColor(0x6B,0x66,0x60); LINE=RGBColor(0xBB,0xB6,0xAE)
SERIF="Georgia"; SANS="Calibri"

d=Document()
for s in d.sections:
    s.top_margin=Inches(0.6); s.bottom_margin=Inches(0.55)
    s.left_margin=Inches(0.85); s.right_margin=Inches(0.85)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hexc); tcPr.append(sh)

def para(text,font=SANS,size=10.5,color=SLATE,bold=False,ital=False,
         align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.06):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    r=p.add_run(text); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=color; f.bold=bold; f.italic=ital
    return p

def writeline(prompt):
    para(prompt,SANS,10,SLATE,bold=True,before=4,after=1)
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    r=p.add_run("_"*94); r.font.color.rgb=LINE; r.font.size=Pt(10)

# Header
para("PUBLICLOGIC",SANS,11,GOLD,bold=True,after=0)
para("Proof & Case Example",SERIF,24,SLATE,bold=True,before=0,after=0)
para("How we prove the work — and how we turn each engagement into a quantified, nameable case.",
     SANS,11,GREEN,ital=True,before=1,after=5)

para("FIVE KINDS OF PROOF — AND THE NUMBER TO CAPTURE",SANS,10,GREEN,bold=True,after=2)
tbl=d.add_table(rows=0,cols=3); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.autofit=False
hdr=tbl.add_row().cells
for i,t in enumerate(["Proof","What it shows","The metric to capture"]):
    hdr[i].width=Inches([1.4,3.0,2.5][i]); shade(hdr[i],"1A1D20")
    r=hdr[i].paragraphs[0].add_run(t); r.font.name=SANS; r.font.size=Pt(9.5); r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
for a,b,c in [
 ("Continuity","A function survives when the person who ran it leaves.","Continuity/transfer risks identified in N functions; successor-readiness from X to Y."),
 ("Record","Every claim, number, and decision traces to a source.","Records consolidated into one governed record; % of decisions now traceable."),
 ("Readiness","The organization can absorb and sustain the change.","Readiness gaps found before the build; # sequenced ahead of spend."),
 ("Adoption","Staff actually use the system, not just receive it.","% of the team using it at 30/90 days; tasks no longer routed through one person."),
 ("Outcome","The work moved forward and got funded.","Permit timeline reduced by X; grant applications supported; $ secured; project delivered."),
]:
    row=tbl.add_row().cells
    row[0].width=Inches(1.4); row[1].width=Inches(3.0); row[2].width=Inches(2.5)
    shade(row[0],"0F4C3A")
    r0=row[0].paragraphs[0].add_run(a); r0.font.name=SANS; r0.font.size=Pt(10); r0.font.bold=True; r0.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    r1=row[1].paragraphs[0].add_run(b); r1.font.name=SANS; r1.font.size=Pt(9.5); r1.font.color.rgb=SLATE
    r2=row[2].paragraphs[0].add_run(c); r2.font.name=SANS; r2.font.size=Pt(9.5); r2.font.color.rgb=GREEN

para("Honest note: PublicLogic is early. Until a client approves a named case, we speak about prior work "
     "in general terms. This sheet is how we stop doing that — every engagement fills in the template "
     "below, so 'similar work has helped…' becomes a specific, quantified, nameable result.",
     SANS,9.5,MUTE,ital=True,before=4,after=6)

# Case template
para("CASE EXAMPLE — one page per engagement (fill in as we go)",SANS,10,GREEN,bold=True,after=1)
writeline("Client / function (name once approved; “a Central MA town” until then):")
writeline("The situation — what was stuck, fragile, or at risk:")
writeline("The trigger — turnover, a deadline, a stalled project, a single point of failure:")
writeline("What we did — the offer and the work (Map / Scan / Sprint / Implementation):")
para("What changed — the numbers (use the metrics above):",SANS,10,SLATE,bold=True,before=4,after=1)
for m in ["Continuity:  ____ risks identified · successor-readiness ____ → ____",
          "Record:  ____ sources consolidated · ____% of decisions now traceable",
          "Readiness:  ____ gaps found and sequenced before spend",
          "Adoption:  ____% of the team using it at 90 days",
          "Outcome:  permit/timeline ____ · grants supported ____ · $ ____ · delivered? ____"]:
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(1)
    r=p.add_run("☐  "+m); r.font.name=SANS; r.font.size=Pt(10); r.font.color.rgb=SLATE
writeline("Proof artifact on file (workbook, register, successor test, funded application):")
writeline("What the client said (one quote, with permission):")

para("Two tests for a real case:  (1) the number is specific and sourced;  (2) the client would let us "
     "say it. If both aren't true yet, keep it general — credibility is worth more than a claim.",
     SANS,9.5,GREEN,ital=True,before=5,after=5)

runs=para  # alias not used
para("The point isn't to be needed forever. It's to make the next person's job easier.",
     SERIF,12,GREEN,bold=True,ital=True,align=WD_ALIGN_PARAGRAPH.CENTER,after=1)
para("PublicLogic LLC  ·  Private & Confidential",SANS,8.5,MUTE,align=WD_ALIGN_PARAGRAPH.CENTER,before=0,after=0)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"final_deck","PublicLogic - Proof & Case Example.docx")
os.makedirs(os.path.dirname(out),exist_ok=True)
d.save(out); print("wrote",out)
