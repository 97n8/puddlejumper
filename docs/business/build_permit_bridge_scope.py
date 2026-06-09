#!/usr/bin/env python3
"""
PublicLogic — Permit & Bridge Sprint: sellable one-page Scope & Fee Sheet.
Closes the "what exactly am I buying?" gap. White-glove navigation for a
project sponsor with a specific stuck/starting project. Mirrors the Map +
Scan sheets. Editable .docx.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SLATE=RGBColor(0x1A,0x1D,0x20); GREEN=RGBColor(0x0F,0x4C,0x3A)
GOLD=RGBColor(0xA9,0x77,0x2F); MUTE=RGBColor(0x6B,0x66,0x60)
SERIF="Georgia"; SANS="Calibri"

d=Document()
for s in d.sections:
    s.top_margin=Inches(0.55); s.bottom_margin=Inches(0.5)
    s.left_margin=Inches(0.8); s.right_margin=Inches(0.8)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hexc); tcPr.append(sh)

def para(text,font=SANS,size=10.5,color=SLATE,bold=False,ital=False,
         align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.05):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    r=p.add_run(text); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=color; f.bold=bold; f.italic=ital
    return p

def runs(p_runs,size=10.5,align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.05):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    for (t,font,col,bold,ital) in p_runs:
        r=p.add_run(t); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=col; f.bold=bold; f.italic=ital
    return p

# Header
para("PUBLICLOGIC",SANS,11,GOLD,bold=True,after=0)
para("Permit & Bridge Sprint",SERIF,24,SLATE,bold=True,before=0,after=0)
para("Scope & Fee Sheet  ·  White-glove help moving a project through the public system.",SANS,11,GREEN,ital=True,before=1,after=5)

para("WHAT IT IS",SANS,10,GREEN,bold=True,after=1)
para("You have a specific project — a building, a use, a program, a development — that has to move "
     "through permits, boards, funding, and stakeholders. Permit & Bridge is the engagement where we "
     "carry it through that public system with you, so it keeps moving and nothing falls through the "
     "cracks between the people involved.",after=5)

para("WHO IT'S FOR",SANS,10,GREEN,bold=True,after=1)
para("Project sponsors — developers, towns, nonprofits, and businesses — with one real project that is "
     "stalled, about to start, or stuck between offices. (If you only need to know which path you're on, "
     "start with a $500 Permit Path Scan instead.)",SANS,9.5,MUTE,ital=True,after=5)

para("WHAT YOU GET",SANS,10,GREEN,bold=True,after=2)
tbl=d.add_table(rows=0,cols=2); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.autofit=False
for label,desc in [
 ("Approvals Map","Every permit, board, and approval the project needs — in the order they have to happen."),
 ("Stakeholder & Board Plan","Who must say yes, in what sequence, and what each one will want to see."),
 ("Submission Checklist","The documents, drawings, and filings each body will require — gathered before you go."),
 ("Coordination & Hearing Support","We manage the path, prep the submissions, and ready you for the meetings and hearings."),
 ("Funding Alignment (optional)","Line the approvals up with grant or funding timing so nothing expires waiting."),
 ("Handoff Record","A clear record of how you got there and who owns what next — so it survives the project."),
]:
    row=tbl.add_row().cells
    row[0].width=Inches(2.25); row[1].width=Inches(4.65)
    shade(row[0],"0F4C3A")
    p0=row[0].paragraphs[0]; r0=p0.add_run(label); r0.font.name=SANS; r0.font.size=Pt(10); r0.font.bold=True; r0.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    p1=row[1].paragraphs[0]; r1=p1.add_run(desc); r1.font.name=SANS; r1.font.size=Pt(10); r1.font.color.rgb=SLATE

para("FEE",SANS,10,GREEN,bold=True,before=7,after=1)
runs([("$7,500 – $15,000",SERIF,SLATE,True,False),
      ("   fixed sprint fee, scoped to the number of approvals and bodies involved.",SANS,MUTE,False,True)],size=12,after=1)
para("Non-contingent — our fee never depends on an approval being granted. Many sponsors continue into "
     "White-Glove Implementation ($3,500–$8,500/month) to carry the path to the finish. A prior Permit "
     "Path Scan or Stewardship Map credits toward this sprint; credits don't stack.",SANS,9.5,MUTE,ital=True,after=5)

para("WHY IT MATTERS",SANS,10,GREEN,bold=True,after=1)
runs([("Most sponsors understand their project. What stalls them is the public system around it. ",SANS,SLATE,True,False),
      ("We become the bridge — translating between the project and the system so good projects don't die "
       "in the gaps between offices.",SANS,SLATE,False,False)],after=4)

para("THE PROMISE",SANS,10,GREEN,bold=True,after=1)
runs([("You leave with a clear path forward, a clear record of how you got there, and a clear "
       "understanding of who owns what next.",SERIF,SLATE,True,True)],size=12,after=5)

para("PROOF",SANS,10,GREEN,bold=True,after=1)
para("Similar reviews have helped project sponsors understand permitting, stakeholder, and board "
     "pathways before committing significant time and money.",SANS,10,SLATE,ital=True,after=6)

runs([("Permit & Bridge helps the public understand the path and helps project sponsors move through it.",SERIF,GREEN,True,True)],
     size=12,align=WD_ALIGN_PARAGRAPH.CENTER,after=2)
para("PublicLogic does not replace the municipality, inspector, planner, or permitting authority. © PublicLogic LLC.",
     SANS,8.5,MUTE,ital=True,align=WD_ALIGN_PARAGRAPH.CENTER,after=1)
runs([("PublicLogic LLC  ·  Nathan Boudreau, MPA / MCPPO  ·  nate@publiclogic.org  ·  Private & Confidential",SANS,MUTE,False,False)],
     size=8.5,align=WD_ALIGN_PARAGRAPH.CENTER,before=0,after=0)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"final_deck","PublicLogic - Permit & Bridge Sprint - Scope & Fee Sheet.docx")
os.makedirs(os.path.dirname(out),exist_ok=True)
d.save(out); print("wrote",out)
