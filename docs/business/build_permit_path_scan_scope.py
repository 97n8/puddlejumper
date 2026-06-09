#!/usr/bin/env python3
"""
PublicLogic — Permit Path Scan: sellable one-page Scope & Fee Sheet.
Tier 1 on the offer ladder: cheap triage between the free LogicCommons tools
and the paid Stewardship Map. Answers "what path am I probably on?"
Editable .docx, on-brand, honest. Reuses the Map sheet's look.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SLATE=RGBColor(0x1A,0x1D,0x20); GREEN=RGBColor(0x0F,0x4C,0x3A)
GOLD=RGBColor(0xA9,0x77,0x2F); MUTE=RGBColor(0x6B,0x66,0x60)
SERIF="Georgia"; SANS="Calibri"

d=Document()
for s in d.sections:
    s.top_margin=Inches(0.6); s.bottom_margin=Inches(0.55)
    s.left_margin=Inches(0.8); s.right_margin=Inches(0.8)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hexc); tcPr.append(sh)

def para(text,font=SANS,size=10.5,color=SLATE,bold=False,ital=False,
         align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.05):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    r=p.add_run(text); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=color; f.bold=bold; f.italic=ital
    return p

def runs(p_runs,size=10.5,align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.05):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    for (t,font,col,bold,ital) in p_runs:
        r=p.add_run(t); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=col; f.bold=bold; f.italic=ital
    return p

# ---- Header
para("PUBLICLOGIC",SANS,11,GOLD,bold=True,after=0)
para("Permit Path Scan",SERIF,24,SLATE,bold=True,before=0,after=0)
para("Scope & Fee Sheet  ·  Cheap triage — “what path am I probably on?”",SANS,11,GREEN,ital=True,before=1,after=6)

# ---- Where it sits
para("WHERE IT SITS",SANS,10,GREEN,bold=True,after=1)
para("Between the free LogicCommons tools and a full paid diagnostic. The free templates help you "
     "understand the path in general. The Permit Path Scan tells YOU, for YOUR specific question, the "
     "path you're most likely on — who to ask first, and what to expect — before you spend real money or time.",
     after=6)

# ---- What it answers
para("WHAT IT ANSWERS",SANS,10,GREEN,bold=True,after=1)
for t in [
 "Is this a zoning, building, health, conservation, or licensing question — or several at once?",
 "Which board or office actually handles it, and who do I talk to first?",
 "Do I likely need a permit, a variance, a special permit, or nothing at all?",
 "What documents will I probably be asked for?",
 "Roughly how long does this path usually take, and where do people get stuck?",
]:
    runs([("—  ",SANS,GOLD,True,False),(t,SANS,SLATE,False,False)],after=1)

# ---- What you get
para("WHAT YOU GET",SANS,10,GREEN,bold=True,before=6,after=2)
tbl=d.add_table(rows=0,cols=2); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
tbl.autofit=False
for label,desc in [
 ("Likely Path","The route you're probably on, in plain language."),
 ("First Contact","The specific board, office, or person to start with."),
 ("Document Checklist","What you'll most likely be asked to bring."),
 ("Watch-outs","The common places this kind of request stalls."),
]:
    row=tbl.add_row().cells
    row[0].width=Inches(1.9); row[1].width=Inches(5.0)
    shade(row[0],"0F4C3A")
    p0=row[0].paragraphs[0]; r0=p0.add_run(label); r0.font.name=SANS; r0.font.size=Pt(10); r0.font.bold=True; r0.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    p1=row[1].paragraphs[0]; r1=p1.add_run(desc); r1.font.name=SANS; r1.font.size=Pt(10); r1.font.color.rgb=SLATE

# ---- Fee
para("FEE",SANS,10,GREEN,bold=True,before=8,after=1)
runs([("$250 – $750",SERIF,SLATE,True,False),
      ("   flat fee, depending on how many systems the question touches.",SANS,MUTE,False,True)],size=12,after=1)
para("Turnaround: a few business days. The Scan is triage, not a guarantee — it points you at the most "
     "likely path so you can move with confidence; the permitting authority always makes the final call. "
     "If you go on to a Stewardship Map, the Scan fee credits toward it.",SANS,9.5,MUTE,ital=True,after=6)

# ---- What it is NOT
para("WHAT IT IS NOT",SANS,10,GREEN,bold=True,after=1)
runs([("It does not replace the municipality, inspector, planner, or permitting authority. ",SANS,SLATE,True,False),
      ("It helps you ask better questions, gather the right information, and avoid wasting time and money "
       "on the wrong path.",SANS,SLATE,False,False)],after=8)

# ---- Closing
runs([("LogicCommons helps people start.  PublicLogic helps them carry it through.",SERIF,GREEN,True,True)],
     size=12,align=WD_ALIGN_PARAGRAPH.CENTER,after=2)
runs([("Implementation is the service.  Institutional Stewardship is the method.",SANS,GOLD,True,False)],
     size=9.5,align=WD_ALIGN_PARAGRAPH.CENTER,after=8)
runs([("PublicLogic LLC",SANS,SLATE,True,False),
      ("   ·   Nathan Boudreau, MPA / MCPPO   ·   Dr. Allison Weiss Rothschild   ·   Private & Confidential",
       SANS,MUTE,False,False)],size=8.5,align=WD_ALIGN_PARAGRAPH.CENTER,before=0,after=0)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"final_deck","PublicLogic - Permit Path Scan - Scope & Fee Sheet.docx")
os.makedirs(os.path.dirname(out),exist_ok=True)
d.save(out); print("wrote",out)
