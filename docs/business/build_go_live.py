#!/usr/bin/env python3
"""
PublicLogic v1 — START HERE & Go-Live Checklist.
The front page of the package: what's inside, what to send whom, and the
exact steps to go live this week. Editable .docx.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SLATE=RGBColor(0x1A,0x1D,0x20); GREEN=RGBColor(0x0F,0x4C,0x3A)
GOLD=RGBColor(0xA9,0x77,0x2F); MUTE=RGBColor(0x6B,0x66,0x60); RED=RGBColor(0x8B,0x2A,0x1A)
SERIF="Georgia"; SANS="Calibri"

d=Document()
for s in d.sections:
    s.top_margin=Inches(0.6); s.bottom_margin=Inches(0.55)
    s.left_margin=Inches(0.85); s.right_margin=Inches(0.85)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hexc); tcPr.append(sh)
def para(text,font=SANS,size=10.5,color=SLATE,bold=False,ital=False,
         align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.06):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    r=p.add_run(text); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=color; f.bold=bold; f.italic=ital
    return p
def runs(p_runs,size=10.5,before=2,after=2,sp=1.06):
    p=d.add_paragraph(); pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    for (t,font,col,bold,ital) in p_runs:
        r=p.add_run(t); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=col; f.bold=bold; f.italic=ital
    return p
def h(text): para(text,SANS,11,GREEN,bold=True,before=9,after=2)
def step(n,text):
    runs([(f"  {n}   ",SANS,RGBColor(0xFF,0xFF,0xFF),True,False)],before=2,after=1)  # placeholder; replaced below
def table(headers,rows,widths,headfill="1A1D20",firstbold=True):
    t=d.add_table(rows=0,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    if any(headers):
        hr=t.add_row().cells
        for i,htxt in enumerate(headers):
            hr[i].width=Inches(widths[i]); shade(hr[i],headfill)
            r=hr[i].paragraphs[0].add_run(htxt); r.font.name=SANS; r.font.size=Pt(9.5); r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].width=Inches(widths[i])
            r=cells[i].paragraphs[0].add_run(val); r.font.name=SANS; r.font.size=Pt(10); r.font.color.rgb=SLATE
            if i==0 and firstbold: r.font.bold=True
    return t
def checkstep(text):
    runs([("☐  ",SANS,GOLD,True,False),(text,SANS,SLATE,False,False)],after=2)

# Header
para("PUBLICLOGIC  ·  v1",SANS,10,GOLD,bold=True,after=0)
para("Start Here & Go-Live Checklist",SERIF,24,SLATE,bold=True,before=0,after=0)
para("The front page of the package. What's inside, what to send whom, and how to go live this week.  "
     "·  June 2026  ·  Private & Confidential.",SANS,10,GREEN,ital=True,before=1,after=5)

runs([("The one line:  ",SANS,GREEN,True,False),
      ("Help people understand the path. Help projects move through the path. Leave the path easier for "
       "the next person.",SERIF,SLATE,False,True)],size=12,after=4)

# What's inside
h("WHAT'S IN THIS PACKAGE (13 documents)")
table(["Document","Use it to…"],
 [["Capabilities Workbook v2.0 (.xlsx)","Run the business. Internal strategy, 12 tabs. Start here."],
  ["Move Forward (.docx)","Brief Allie / a partner in two pages."],
  ["The Path — one-page visual (.pptx)","Show the whole model in one picture."],
  ["One Pager (.pptx)","Send a fast overview."],
  ["Capabilities & Positioning deck (.pptx)","Walk a meeting through the story."],
  ["Stewardship Map — Scope & Fee (.docx)","Sell the $2,500–$7,500 entry diagnostic."],
  ["Permit Path Scan — Scope & Fee (.docx)","Sell the $500 flat triage."],
  ["Permit & Bridge Sprint — Scope & Fee (.docx)","Sell the $7,500–$15,000 navigation sprint."],
  ["Permit Path Scan — 30-Day Cash Offer (.docx)","Generate cash this month (the run book)."],
  ["LogicCommons — Can I Do This? (.docx)","Hand out free; it feeds the Scan."],
  ["Five Outreach Emails (.docx)","Start conversations with warm contacts."],
  ["Proof & Case Example (.docx)","Turn each engagement into a quantified case."],
  ["Website Copy (.docx)","Reference for the live site (already built in apps/web)."]],
 [3.3,3.5])

# What to send whom
h("WHAT TO SEND WHOM")
table(["When someone is…","Send / sell"],
 [["A resident or owner stuck on a permit","The free “Can I Do This?” worksheet → then a $500 Permit Path Scan."],
  ["A warm municipal contact","The Stewardship Map Scope & Fee sheet ($2,500–$7,500)."],
  ["A developer / sponsor with a stuck project","The Permit & Bridge Sprint Scope & Fee sheet."],
  ["Anyone in your network","One of the Five Outreach Emails — ask for one conversation."],
  ["A closed engagement","Fill in the Proof & Case Example before you forget the numbers."]],
 [3.1,3.7])

# Go-live checklist
h("GO LIVE THIS WEEK — CHECKLIST")
checkstep("Domain & email: confirm publiclogic.org; stand up the nate@publiclogic.org mailbox.")
checkstep("Deploy the site: apps/web (Next.js 15). It builds clean — deploy to Vercel and point the domain at it.")
checkstep("Wire the contact form: set CONTACT_TO = nate@publiclogic.org plus either RESEND_API_KEY + CONTACT_FROM (a verified sender) OR CONTACT_WEBHOOK_URL. Until then, submissions are saved to the server log. (See ENV_REFERENCE.md.)")
checkstep("Confirm the worksheet downloads: /logiccommons/can-i-do-this (Download .docx + Print / Save as PDF).")
checkstep("Pricing is locked: Permit Path Scan $500 flat · Rush $750 · Complex $1,250 (E&O) · Follow-up $150 · Stewardship Map $2,500–$7,500.")
runs([("E&O before the $1,250 tier: ",SANS,RED,True,False),
      ("confirm professional-liability coverage before taking a developer's complex scan, and keep every "
       "scan in “likely / probably.”",SANS,SLATE,False,True)],size=9.5,after=2)
checkstep("Start selling: send the five outreach emails and run the 30-Day Cash Offer — mine the network as a referral engine, not a cold list.")

# 30-day money
h("THE 30-DAY MONEY PLAN")
table(["Scenario","Cash"],
 [["6 scans × $500","$3,000"],["10 scans × $500","$5,000"],["8 scans + 2 complex","$6,500+"]],
 [3.6,2.0])
runs([("The fastest cash path isn't convincing towns to buy continuity. ",SANS,SLATE,True,False),
      ("It's helping people with real projects stop bleeding time in permitting confusion. The Scan is the "
       "wedge and the proof — not the ceiling.",SANS,GREEN,False,True)],size=10,after=4)

para("PublicLogic LLC  ·  Nathan Boudreau, MPA / MCPPO  ·  Dr. Allison Weiss Rothschild  ·  "
     "nate@publiclogic.org  ·  978-807-0829  ·  Private & Confidential",
     SANS,8.5,MUTE,align=WD_ALIGN_PARAGRAPH.CENTER,before=4,after=0)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"final_deck","PublicLogic - START HERE & Go-Live Checklist.docx")
os.makedirs(os.path.dirname(out),exist_ok=True)
d.save(out); print("wrote",out)
