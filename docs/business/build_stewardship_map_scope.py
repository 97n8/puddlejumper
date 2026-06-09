#!/usr/bin/env python3
"""
PublicLogic — Stewardship Map: sellable one-page Scope & Fee Sheet.
A real document Nathan can send a Tier-1 prospect. Editable .docx, on-brand,
honest. The Map is the safe first engagement (Tier 2 on the offer ladder).
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SLATE=RGBColor(0x1A,0x1D,0x20); GREEN=RGBColor(0x0F,0x4C,0x3A)
GOLD=RGBColor(0xA9,0x77,0x2F); MUTE=RGBColor(0x6B,0x66,0x60)
SERIF="Georgia"; SANS="Calibri"

d=Document()
for s in d.sections:
    s.top_margin=Inches(0.6); s.bottom_margin=Inches(0.55)
    s.left_margin=Inches(0.8); s.right_margin=Inches(0.8)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hexc); tcPr.append(sh)

def para(text,font=SANS,size=10.5,color=SLATE,bold=False,ital=False,
         align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.05):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    r=p.add_run(text); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=color; f.bold=bold; f.italic=ital
    return p

def runs(p_runs,size=10.5,align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.05):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    for (t,font,col,bold,ital) in p_runs:
        r=p.add_run(t); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=col; f.bold=bold; f.italic=ital
    return p

# ---- Header
para("PUBLICLOGIC",SANS,11,GOLD,bold=True,after=0)
para("Stewardship Map",SERIF,24,SLATE,bold=True,before=0,after=0)
para("Scope & Fee Sheet  ·  The safe first engagement",SANS,11,GREEN,ital=True,before=1,after=6)

# ---- What it is
para("WHAT IT IS",SANS,10,GREEN,bold=True,after=1)
para("Before anyone commits real money to a project, a build, or a hire, the Stewardship Map shows "
     "how work, knowledge, authority, records, and risk actually move through your organization — "
     "and where the next person would get stuck. It is a focused, fixed-fee diagnostic. It is the "
     "safest way to start working together, because it tells both of us exactly what should come next.",
     after=6)

# ---- What we look at
para("WHAT WE LOOK AT",SANS,10,GREEN,bold=True,after=1)
for t in [
 "How the work actually flows — not the org chart, the real path.",
 "Where knowledge lives, and what leaves the building when a person does.",
 "Who holds authority, and where decisions stall or bottleneck.",
 "How records are kept, and whether a claim or number can be traced to a source.",
 "Where funding, permitting, or projects are exposed to a single point of failure.",
]:
    runs([("—  ",SANS,GOLD,True,False),(t,SANS,SLATE,False,False)],after=1)

# ---- What you get
para("WHAT YOU GET",SANS,10,GREEN,bold=True,before=6,after=2)
tbl=d.add_table(rows=0,cols=2); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
tbl.autofit=False
for label,desc in [
 ("Stewardship Map","A clear picture of how work, knowledge, authority, and risk move today."),
 ("Continuity Risks","The points where turnover or a gap would actually break something."),
 ("Readiness Findings","An honest read on what your organization can absorb and sustain right now."),
 ("Priority Roadmap","A short, sequenced list of the next moves — with what each one is worth."),
]:
    row=tbl.add_row().cells
    row[0].width=Inches(1.9); row[1].width=Inches(5.0)
    shade(row[0],"0F4C3A")
    p0=row[0].paragraphs[0]; r0=p0.add_run(label); r0.font.name=SANS; r0.font.size=Pt(10); r0.font.bold=True; r0.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    p1=row[1].paragraphs[0]; r1=p1.add_run(desc); r1.font.name=SANS; r1.font.size=Pt(10); r1.font.color.rgb=SLATE

# ---- Fee
para("FEE",SANS,10,GREEN,bold=True,before=8,after=1)
runs([("$2,500 – $7,500",SERIF,SLATE,True,False),
      ("   fixed fee, scoped to the size of the organization and the question.",SANS,MUTE,False,True)],size=12,after=1)
para("Typical engagement: 2–4 weeks. Fixed-fee and non-contingent — our fee never depends on a "
     "funding outcome or a decision going a particular way.",SANS,9.5,MUTE,ital=True,after=1)
runs([("Credit policy:  ",SANS,GREEN,True,False),
      ("if you continue, the Map fee credits toward the next engagement. Credits don't stack — "
       "up to one tier's fee credits forward, so you never pay twice for the same step.",
       SANS,MUTE,False,True)],size=9.5,after=6)

# ---- Why start here
para("WHY START HERE",SANS,10,GREEN,bold=True,after=1)
runs([("Low risk. ",SANS,SLATE,True,False),
      ("Small, fixed, and bounded — you learn what's really going on before committing to anything bigger. "
       "Most clients use the Map to decide what to do next, and roughly half of what it finds is something "
       "they hadn't seen coming.",SANS,SLATE,False,False)],after=8)

# ---- Closing line + contact
runs([("Every system should make it easier for the next person to do the right thing.",SERIF,GREEN,True,True)],
     size=12,align=WD_ALIGN_PARAGRAPH.CENTER,after=2)
runs([("Implementation is the service.  Institutional Stewardship is the method.",SANS,GOLD,True,False)],
     size=9.5,align=WD_ALIGN_PARAGRAPH.CENTER,after=8)
runs([("PublicLogic LLC",SANS,SLATE,True,False),
      ("   ·   Nathan Boudreau, MPA / MCPPO   ·   Dr. Allison Weiss Rothschild   ·   Private & Confidential",
       SANS,MUTE,False,False)],size=8.5,align=WD_ALIGN_PARAGRAPH.CENTER,before=0,after=0)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"final_deck","PublicLogic - Stewardship Map - Scope & Fee Sheet.docx")
os.makedirs(os.path.dirname(out),exist_ok=True)
d.save(out); print("wrote",out)
