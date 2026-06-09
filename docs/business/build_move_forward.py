#!/usr/bin/env python3
"""
PublicLogic — Move Forward: v1 Business Stack (cover + index + 90-day plan).
The single front-door document for the v1 stack. On-brand, honest, short.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SLATE=RGBColor(0x1A,0x1D,0x20); GREEN=RGBColor(0x0F,0x4C,0x3A)
GOLD=RGBColor(0xA9,0x77,0x2F); MUTE=RGBColor(0x6B,0x66,0x60)
SERIF="Georgia"; SANS="Calibri"

d=Document()
for s in d.sections:
    s.top_margin=Inches(0.6); s.bottom_margin=Inches(0.55)
    s.left_margin=Inches(0.8); s.right_margin=Inches(0.8)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hexc); tcPr.append(sh)

def para(text,font=SANS,size=10.5,color=SLATE,bold=False,ital=False,
         align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.05):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    r=p.add_run(text); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=color; f.bold=bold; f.italic=ital
    return p

def runs(p_runs,size=10.5,align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.05):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    for (t,font,col,bold,ital) in p_runs:
        r=p.add_run(t); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=col; f.bold=bold; f.italic=ital
    return p

def tbl2(rows,hdr=("","" ),wlabel=2.2,wbody=4.7,head_fill="0F4C3A"):
    t=d.add_table(rows=0,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    if hdr[0] or hdr[1]:
        hr=t.add_row().cells
        for i,(c,txt) in enumerate(zip(hr,hdr)):
            c.width=Inches(wlabel if i==0 else wbody); shade(c,"1A1D20")
            p=c.paragraphs[0]; r=p.add_run(txt); r.font.name=SANS; r.font.size=Pt(9.5); r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    for label,desc,strong in rows:
        row=t.add_row().cells
        row[0].width=Inches(wlabel); row[1].width=Inches(wbody)
        if strong: shade(row[0],head_fill)
        p0=row[0].paragraphs[0]; r0=p0.add_run(label); r0.font.name=SANS; r0.font.size=Pt(10); r0.font.bold=True
        r0.font.color.rgb=(RGBColor(0xFF,0xFF,0xFF) if strong else SLATE)
        p1=row[1].paragraphs[0]; r1=p1.add_run(desc); r1.font.name=SANS; r1.font.size=Pt(10); r1.font.color.rgb=SLATE
    return t

# ---- Header
para("PUBLICLOGIC",SANS,11,GOLD,bold=True,after=0)
para("Move Forward — v1 Business Stack",SERIF,23,SLATE,bold=True,before=0,after=0)
para("The front door. What we are, what's in the stack, and the next 90 days.  ·  June 2026  ·  Private & Confidential",
     SANS,10,GREEN,ital=True,before=1,after=6)

# ---- The company in four lines
para("THE COMPANY, IN FOUR LINES",SANS,10,GREEN,bold=True,after=2)
for t,c in [
 ("Every system should make it easier for the next person to do the right thing.",GREEN),
 ("PublicLogic helps projects move through public systems, and helps organizations keep the capacity to sustain them.",SLATE),
 ("Implementation is the service. Institutional Stewardship is the method.",GOLD),
 ("Good Work + Stewardship = Work That Lasts.",SLATE),
]:
    runs([("— ",SANS,GOLD,True,False),(t,SERIF if c==GREEN else SANS,c,c in (GREEN,GOLD),c==GREEN)],size=11,after=1)

# ---- What breaks (the demand statement)
para("WHAT BREAKS  (why this is needed)",SANS,10,GREEN,bold=True,before=8,after=2)
tbl2([
 ("Turnover","Knowledge leaves with the person.",False),
 ("Unclear ownership","Projects stall — no one holds the next step.",False),
 ("Funding without readiness","Momentum collapses; the window closes.",False),
 ("No continuity system","Work restarts from zero.",False),
 ("Plans without implementation","Reports sit on shelves.",False),
],hdr=("What breaks","What happens"))
runs([("PublicLogic exists because these problems are common, expensive, and preventable.",
       SANS,GREEN,True,True)],size=11,before=3,after=6)

# ---- The architecture / progression
para("THE PROGRESSION  (this is how people actually buy)",SANS,10,GREEN,bold=True,before=8,after=2)
tbl2([
 ("LogicCommons · Understand","“I need help understanding.”  — free public templates, checklists, frameworks.",True),
 ("Permit & Bridge · Navigate","“I need help navigating.”  — find the path through permits, boards, funding.",True),
 ("Stewardship Map · Diagnose","“I need to know what actually needs to happen.”  — the safe paid first step.",True),
 ("PublicLogic · Deliver","“I need help getting this done.”  — sprints, implementation, capacity.",True),
],hdr=("Layer · Verb","What the buyer is saying"))
para("Delivered through Continuity & Stewardship Systems — the records, processes, templates, environments, "
     "accountability, and proof that help the work survive turnover. It's the plumbing, not a product.",
     SANS,9.5,MUTE,ital=True,before=3,after=3)
para("What we're becoming:  LogicCommons = public infrastructure · Permit & Bridge = navigation infrastructure · "
     "PublicLogic = professional services · Continuity & Stewardship Systems = delivery infrastructure.",
     SANS,9.5,MUTE,ital=True,before=0,after=4)
runs([("The one-line version:  ",SANS,GREEN,True,False),
      ("Help people understand the path. Help projects move through the path. Leave the path easier for the next person.",
       SERIF,SLATE,False,True)],size=11.5,before=2,after=6)

# ---- What's in the stack
para("WHAT'S IN THE STACK",SANS,10,GREEN,bold=True,after=2)
tbl2([
 ("Capabilities Workbook v2.0","The operating document. 11 tabs: why we exist, what breaks, what we do, the offer stack, Permit & Bridge, proof, Nathan + Allie, transfer, who we pursue, next 12 months. Internal strategy — start here.",False),
 ("One Pager","The whole thing on one page. For a fast read or a first send.",False),
 ("Capabilities & Positioning deck","The narrative version, for a meeting or a screen-share.",False),
 ("Stewardship Map — Scope & Fee Sheet","Sellable. $2,500–$7,500, fixed, non-contingent. The first real sale.",False),
 ("Permit Path Scan — Scope & Fee Sheet","Sellable. $250–$750 triage. The cheap entry that feeds the Map.",False),
],hdr=("File","What it's for"))

# ---- The offer ladder (quick reference)
para("THE OFFER LADDER  (quick reference)",SANS,10,GREEN,bold=True,before=8,after=2)
tbl2([
 ("Tier 0 — Public Permit Helper","Free / very low cost",False),
 ("Tier 1 — Permit Path Scan","$250 – $750",False),
 ("Tier 2 — Stewardship Map","$2,500 – $7,500",False),
 ("Tier 3 — Permit & Bridge Sprint","$7,500 – $15,000",False),
 ("Tier 4 — White-Glove Implementation","$3,500 – $8,500 / month",False),
 ("Tier 5 — Funding / Grant Build","$5,000 – $25,000 (scope-dependent)",False),
],hdr=("Tier","Price"))
para("Credit policy: a paid tier credits toward the next engagement, but credits don't stack — up to one "
     "tier's fee credits forward. Low entry risk, protected margin.",SANS,9.5,MUTE,ital=True,before=3,after=6)

# ---- The next 90 days
para("THE NEXT 90 DAYS  (the ball moves here, not in another workbook)",SANS,10,GREEN,bold=True,after=2)
tbl2([
 ("1 · Show Allie","Walk the workbook. If she says “yes, that's us,” stop editing and start selling.",True),
 ("2 · Stand up the two MVPs","LogicCommons MVP (a handful of free templates) and Permit & Bridge MVP (the public “Can I do this?” helper). Small and real.",True),
 ("3 · Sell 2–4 Maps","Use the Scope & Fee sheet with warm Tier-1 relationships. The Map is the wedge into everything else.",True),
 ("4 · Make the calls","CMRPC. MRPC. Community Paradigm. Warm municipal contacts. One client conversation beats ten workbook revisions.",True),
],hdr=("Move","What it looks like"))

# ---- Who to call first
para("WHO TO CALL FIRST  (Tier 1)",SANS,10,GREEN,bold=True,before=8,after=1)
runs([("Existing municipal network  ·  Regional planning agencies (CMRPC, MRPC)  ·  Engineering firms  ·  "
       "Municipal consulting firms (Community Paradigm).",SANS,SLATE,False,False),
      ("  Warmest relationships, real projects and budgets, shortest sales cycle.",SANS,MUTE,False,True)],after=8)

# ---- Close
runs([("The point isn't to be needed forever. It's to make the next person's job easier.",SERIF,GREEN,True,True)],
     size=12,align=WD_ALIGN_PARAGRAPH.CENTER,after=2)
runs([("If neither PublicLogic gets stronger nor the client gets less dependent, we're doing consulting instead of stewardship.",
       SANS,GOLD,True,False)],size=9.5,align=WD_ALIGN_PARAGRAPH.CENTER,after=8)
runs([("PublicLogic LLC",SANS,SLATE,True,False),
      ("   ·   Nathan Boudreau, MPA / MCPPO   ·   Dr. Allison Weiss Rothschild   ·   Private & Confidential",
       SANS,MUTE,False,False)],size=8.5,align=WD_ALIGN_PARAGRAPH.CENTER,before=0,after=0)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"final_deck","PublicLogic - Move Forward (v1 Business Stack).docx")
os.makedirs(os.path.dirname(out),exist_ok=True)
d.save(out); print("wrote",out)
