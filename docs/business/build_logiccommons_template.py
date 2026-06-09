#!/usr/bin/env python3
"""
LogicCommons — Template #1: "Can I Do This?" A plain-language permit-path worksheet.
The free public artifact. Genuinely useful triage, not a sales doc. Feeds the
Permit Path Scan only at the very end, gently. Fillable .docx.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SLATE=RGBColor(0x1A,0x1D,0x20); GREEN=RGBColor(0x0F,0x4C,0x3A)
GOLD=RGBColor(0xA9,0x77,0x2F); MUTE=RGBColor(0x6B,0x66,0x60); LINE=RGBColor(0xBB,0xB6,0xAE)
SERIF="Georgia"; SANS="Calibri"

d=Document()
for s in d.sections:
    s.top_margin=Inches(0.6); s.bottom_margin=Inches(0.55)
    s.left_margin=Inches(0.85); s.right_margin=Inches(0.85)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hexc); tcPr.append(sh)

def para(text,font=SANS,size=10.5,color=SLATE,bold=False,ital=False,
         align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.06):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    r=p.add_run(text); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=color; f.bold=bold; f.italic=ital
    return p

def runs(p_runs,size=10.5,align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.06):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    for (t,font,col,bold,ital) in p_runs:
        r=p.add_run(t); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=col; f.bold=bold; f.italic=ital
    return p

def section(label):
    para(label,SANS,10.5,GREEN,bold=True,before=9,after=2)

def writeline(prompt,n=1):
    para(prompt,SANS,10,SLATE,after=1)
    for _ in range(n):
        p=d.add_paragraph(); p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(3)
        r=p.add_run("_"*92); r.font.color.rgb=LINE; r.font.size=Pt(10)

def checkrow(items):
    for it in items:
        runs([("☐  ",SANS,GOLD,True,False),(it,SANS,SLATE,False,False)],after=1)

# ---- Header
para("LOGICCOMMONS  ·  FREE PUBLIC WORKSHEET",SANS,10,GOLD,bold=True,after=0)
para("Can I Do This?",SERIF,24,SLATE,bold=True,before=0,after=0)
para("A plain-language worksheet for figuring out the path before you spend time or money.",
     SANS,11,GREEN,ital=True,before=1,after=4)
para("Fill this out before you call town hall. It won't give you an answer — only the town can do that — "
     "but it will help you ask the right office the right question, and show up with the right information.",
     SANS,9.5,MUTE,ital=True,after=2)

# ---- 1. What
section("1 · WHAT ARE YOU TRYING TO DO?")
writeline("Describe it in one or two plain sentences (e.g., “put a shed in my backyard,” “open a small "
          "food business,” “add an apartment over my garage,” “run an event on my property”).",2)

# ---- 2. Where
section("2 · WHERE WOULD IT HAPPEN?")
checkrow(["My own home / yard","A property I rent or manage","A commercial building or storefront",
          "Vacant land","Public or shared space","Not sure yet"])
writeline("Address or rough location (town / neighborhood):")

# ---- 3. Which systems
section("3 · WHICH SYSTEMS MIGHT THIS TOUCH?  (check any that might apply — guessing is fine)")
tbl=d.add_table(rows=0,cols=2); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.autofit=False
for area,clue in [
 ("Zoning","Is this use allowed here? Setbacks, size, how the land can be used."),
 ("Building","Is anything being built, changed, wired, or plumbed?"),
 ("Health","Food, water, septic, animals, or anything that affects public health?"),
 ("Conservation","Wetlands, water, slopes, trees, or protected land nearby?"),
 ("Licensing","Does the activity itself need a license or permit to operate?"),
 ("Historic / other","Historic district, special overlay, or something unusual?"),
]:
    row=tbl.add_row().cells
    row[0].width=Inches(1.7); row[1].width=Inches(5.2)
    shade(row[0],"0F4C3A")
    p0=row[0].paragraphs[0]; r0=p0.add_run("☐  "+area); r0.font.name=SANS; r0.font.size=Pt(10); r0.font.bold=True; r0.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    p1=row[1].paragraphs[0]; r1=p1.add_run(clue); r1.font.name=SANS; r1.font.size=Pt(9.5); r1.font.color.rgb=SLATE

# ---- 4. First questions
section("4 · THE FIVE QUESTIONS TO ASK FIRST")
for q in [
 "Is what I want to do allowed at this location at all?",
 "Which board or office handles it — and who do I talk to first?",
 "Do I need a permit, a variance, a special permit, a license — or nothing?",
 "What documents or drawings will I be asked to bring?",
 "Roughly how long does this usually take, and is there a public hearing?",
]:
    runs([("•  ",SANS,GOLD,True,False),(q,SANS,SLATE,False,False)],after=1)

# ---- 5. Who handles what
section("5 · WHO USUALLY HANDLES WHAT  (a starting guide — every town is a little different)")
tbl2=d.add_table(rows=0,cols=2); tbl2.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl2.autofit=False
hdr=tbl2.add_row().cells
for i,t in enumerate(["If it's about…","…start with"]):
    hdr[i].width=Inches(2.6 if i==0 else 4.3); shade(hdr[i],"1A1D20")
    r=hdr[i].paragraphs[0].add_run(t); r.font.name=SANS; r.font.size=Pt(9.5); r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
for a,b in [
 ("How land can be used / size / setbacks","Zoning / Planning office or Zoning Board"),
 ("Building, wiring, plumbing, structural","Building Department / Inspector"),
 ("Food, septic, water, animals","Board of Health"),
 ("Wetlands, water, protected land","Conservation Commission"),
 ("Running a business or event","Town Clerk / Licensing authority / Select Board"),
 ("Historic or special districts","Historic Commission / Planning"),
]:
    row=tbl2.add_row().cells
    row[0].width=Inches(2.6); row[1].width=Inches(4.3)
    r0=row[0].paragraphs[0].add_run(a); r0.font.name=SANS; r0.font.size=Pt(10); r0.font.color.rgb=SLATE
    r1=row[1].paragraphs[0].add_run(b); r1.font.name=SANS; r1.font.size=Pt(10); r1.font.bold=True; r1.font.color.rgb=GREEN

# ---- 6. Gather
section("6 · WHAT TO GATHER BEFORE YOU GO")
checkrow(["Your address and parcel / lot number","A simple sketch or photo of what you want to do",
          "Rough measurements (size, distance from property lines)","Any past permits or plans you already have",
          "A short written description of the project"])

# ---- 7. Next step
section("7 · YOUR LIKELY NEXT STEP")
writeline("Based on the above, the first office I should call is:")
writeline("The first question I'll ask them is:")

# ---- When to get help
para("WHEN IT'S WORTH GETTING HELP",SANS,10,GREEN,bold=True,before=9,after=1)
runs([("If your project touches several systems at once, involves a hearing, or you've already been bounced "
       "between offices, a ",SANS,SLATE,False,False),
      ("Permit Path Scan",SANS,GREEN,True,False),
      (" ($250–$750) lays out the likely path, the boards involved, and the documents you'll need — so you "
       "stop guessing. It doesn't replace the town; it helps you move with confidence.",SANS,SLATE,False,False)],after=8)

# ---- Footer / canon
runs([("LogicCommons helps people start.  PublicLogic helps them carry it through.",SERIF,GREEN,True,True)],
     size=11.5,align=WD_ALIGN_PARAGRAPH.CENTER,after=2)
para("This worksheet is a free public tool. It does not replace your municipality, inspector, planner, or "
     "permitting authority, and it is not legal advice. Always confirm with the office that has authority.",
     SANS,8.5,MUTE,ital=True,align=WD_ALIGN_PARAGRAPH.CENTER,after=1)
runs([("PublicLogic LLC  ·  LogicCommons  ·  Private & Confidential draft",SANS,MUTE,False,False)],
     size=8.5,align=WD_ALIGN_PARAGRAPH.CENTER,before=0,after=0)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"final_deck","LogicCommons - Can I Do This (permit-path worksheet).docx")
os.makedirs(os.path.dirname(out),exist_ok=True)
d.save(out); print("wrote",out)
