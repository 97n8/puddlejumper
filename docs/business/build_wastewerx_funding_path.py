#!/usr/bin/env python3
"""
PublicLogic — A Funding Path for WasteWerx (STTR one-pager, branded).
Faithful to the approved copy; rendered in the PublicLogic house style.
Prospect-facing; deal-specific (kept out of the v1 stack).
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SLATE=RGBColor(0x1A,0x1D,0x20); GREEN=RGBColor(0x0F,0x4C,0x3A); GOLD=RGBColor(0xA9,0x77,0x2F)
MUTE=RGBColor(0x6B,0x66,0x60); WHITE=RGBColor(0xFF,0xFF,0xFF); SERIF="Georgia"; SANS="Calibri"
d=Document()
for s in d.sections:
    s.top_margin=Inches(0.55); s.bottom_margin=Inches(0.5); s.left_margin=Inches(0.8); s.right_margin=Inches(0.8)
def shade(c,h):
    tc=c._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),h); tc.append(sh)
def para(t,font=SANS,size=10,color=SLATE,bold=False,ital=False,align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.05):
    p=d.add_paragraph(); p.alignment=align; pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    r=p.add_run(t); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=color; f.bold=bold; f.italic=ital; return p
def h(t): para(t,SANS,10.5,GREEN,bold=True,before=8,after=1)
def bullet(lead,rest):
    p=d.add_paragraph(); pf=p.paragraph_format; pf.space_after=Pt(1); pf.line_spacing=1.04
    r=p.add_run("—  "+lead); r.font.name=SANS; r.font.size=Pt(10); r.font.bold=True; r.font.color.rgb=GOLD
    r2=p.add_run(rest); r2.font.name=SANS; r2.font.size=Pt(10); r2.font.color.rgb=SLATE
def greenbox(title,body):
    t=d.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    cell=t.rows[0].cells[0]; cell.width=Inches(6.9); shade(cell,"0F4C3A")
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(2)
    r=p.add_run(title); r.font.name=SANS; r.font.size=Pt(9.5); r.font.bold=True; r.font.color.rgb=RGBColor(0xD8,0xC8,0x9A)
    p2=cell.add_paragraph(); r2=p2.add_run(body); r2.font.name=SERIF; r2.font.size=Pt(11.5); r2.font.color.rgb=WHITE; r2.font.italic=True

# Header
para("PUBLICLOGIC",SANS,10,GOLD,bold=True,after=0)
para("A Funding Path for WasteWerx",SERIF,22,SLATE,bold=True,before=0,after=0)
para("WasteWerx  +  University of Central Florida  +  PublicLogic",SANS,11,GREEN,ital=True,before=1,after=4)

para("The short version. There is federal money built for exactly this situation — a company with working "
     "technology that teams up with a university to prove it out. It’s called STTR, and the partnership "
     "structure it requires is the one we already have in front of us. WasteWerx is the applicant and "
     "management lead; the lead-researcher structure is matched to the target agency’s STTR rules, with UCF "
     "serving as the required research-institution partner — so WasteWerx stays in control. UCF brings the "
     "independent lab credibility that turns “we say it works” into “a research university confirmed it "
     "works.” PublicLogic coordinates grant development, partnership structure, compliance, and project "
     "stewardship across the collaboration. This page lays out who does what and why it’s worth your time.",
     after=5)

greenbox("THE QUESTION THE GRANT PAYS US TO ANSWER",
     "Can the WasteWerx process turn everyday waste — tires, mixed plastics, wood — into clean jet and diesel "
     "fuel that meets the industry’s quality bar (the ASTM standard), in one pass, with a carbon footprint a "
     "buyer can trust? A university saying yes to that is worth far more than a vendor saying it.")
para("That’s a feasibility question — which is precisely what this grant funds. Answer it well and the next, "
     "bigger round of money (and serious buyers) opens up.",before=3,after=5)

h("WHO DOES WHAT")
cols=[("WasteWerx","The technology and the proof",
       ["Run the feedstocks — tires, plastics, wood — through the system",
        "Produce the fuel samples and co-products for testing",
        "Share the run data: yields, uptime, operating conditions",
        "Own the equipment, the IP, and the path to building more",
        "Line up the buyers and licensing interest"]),
      ("UCF","The independent proof",
       ["Test the fuel against the ASTM quality standards",
        "Confirm what the fuel is and how clean it burns",
        "Help shorten the ASTM certification timeline",
        "Measure the real carbon footprint, start to finish",
        "Put a university name on the results — and supply the research lead the grant requires"]),
      ("PublicLogic","The build, the paperwork, the adoption",
       ["Write and assemble the full grant package",
        "Structure the agreement and split the funds correctly",
        "Run all reporting and compliance for the life of the award",
        "Dr. Rothschild leads adoption — people, training, and trust",
        "Stand as the steward holding the three parties together"])]
tb=d.add_table(rows=0,cols=3); tb.alignment=WD_TABLE_ALIGNMENT.CENTER; tb.autofit=False
hr=tb.add_row().cells
for i,(name,tag,_ ) in enumerate(cols):
    hr[i].width=Inches(2.3); shade(hr[i],"1A1D20")
    p=hr[i].paragraphs[0]; r=p.add_run(name); r.font.name=SERIF; r.font.size=Pt(12); r.font.bold=True; r.font.color.rgb=WHITE
    p2=hr[i].add_paragraph(); r2=p2.add_run(tag); r2.font.name=SANS; r2.font.size=Pt(8.5); r2.font.italic=True; r2.font.color.rgb=RGBColor(0xD8,0xC8,0x9A)
br=tb.add_row().cells
for i,(_,_,items) in enumerate(cols):
    br[i].width=Inches(2.3); cell=br[i]; cell.paragraphs[0].text=""
    for j,it in enumerate(items):
        p=cell.paragraphs[0] if j==0 else cell.add_paragraph(); p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.02
        r=p.add_run("• "); r.font.name=SANS; r.font.size=Pt(9); r.font.color.rgb=GOLD
        r2=p.add_run(it); r2.font.name=SANS; r2.font.size=Pt(9); r2.font.color.rgb=SLATE

para("On the money: the large majority of the grant funds the actual research — WasteWerx running the system "
     "and UCF testing the results. PublicLogic is paid as a smaller line for building and managing the award "
     "(grant development, compliance, and stewardship), separate from the research work itself. It is not a "
     "cut of your share or UCF’s.",before=5,after=4)

h("THE ADOPTION LAYER — AND WHO LEADS IT")
para("UCF proves the fuel is real. Getting a town, a hauler, or a plant operator to actually run the system, "
     "trust it, and keep using it is a different problem. PublicLogic’s organizational implementation work is "
     "led by Dr. Allison Weiss Rothschild. She holds a PsyD in Leadership Psychology from William James "
     "College and an MSW from Boston College, with licenses as a Clinical Social Worker (LICSW) and Behavior "
     "Analyst (BCBA/LABA), and has spent her career inside multi-site, regulated human-services "
     "organizations. Her doctoral work looked at why organizations struggle to hold onto new practices and "
     "staff buy-in after a strong start — the same problem a new technology runs into once the demonstration "
     "phase ends. She draws on CFIR, a framework researchers use to assess whether a new system actually "
     "takes hold in an organization, to build the training and adoption plan that carries a project past the "
     "lab bench. Reviewers increasingly look for teams with a credible path to real-world adoption, not just "
     "technical feasibility, and that’s the piece PublicLogic adds to the partnership.",after=4)

h("WHAT THIS ASKS OF YOU")
bullet("Access — ","Let UCF independently test samples from your runs. The results are yours and the partnership’s.")
bullet("Data — ","Share enough run data for the university to do honest measurement. Your IP and reactor design stay yours.")
bullet("Time — ","A modest in-kind commitment of machine time and materials, which also counts toward the grant’s cost-share.")
para("PublicLogic does the writing, the budgeting, the filing, and the reporting. The lift on your side is "
     "the technology you’re already running.",ital=True,color=MUTE,size=9.5,before=2,after=4)

h("WHAT YOU GET")
bullet("Funding — ","Non-dilutive federal money — a grant, not an investment. No equity given up.")
bullet("Credibility — ","A research university’s independent confirmation of your fuel and its carbon footprint — the single most useful thing you can hand a buyer or the next investor.")
bullet("The next step — ","A clear runway to the larger follow-on round and to serious offtake conversations.")
bullet("Staying power — ","A documented adoption plan that makes the whole thing real beyond the lab bench.")

h("THE TIMING")
para("The program has been reauthorized through September 30, 2031, and the next step is matching the project "
     "to the right agency solicitation cycle — for a waste-to-fuel process, most likely DOE or USDA. From "
     "there the work is assembling the package and pairing the project with the right UCF researcher — we’ll "
     "work through UCF’s research office to land that fit, and PublicLogic coordinates the rest. One honest "
     "note: every grant is competitive, and we won’t put anything forward that doesn’t position WasteWerx as "
     "a clear contender. The point of getting in early is to build exactly that kind of application.",after=4)

h("THE SMALL FIRST STEP")
para("To get moving without a big commitment, PublicLogic will run the initial grant-fit and due-diligence "
     "work — confirming the best agency, the solicitation window, and how competitive WasteWerx would be — "
     "for a modest fixed fee of $1,500. If we proceed to the application, it credits toward the "
     "grant-management work. It’s the fastest way to turn this from a good idea into a filed application.",after=4)

greenbox("WHAT WE’D NAIL DOWN TOGETHER",
     "What gets tested — which feedstocks, which fuel fractions, and what “success” looks like.   "
     "How it’s structured — WasteWerx as prime, the UCF researcher, and an agreement that protects your IP.   "
     "Who carries what — confirming the split so the budget and roles are clear before anything is drafted.")

para("That’s the agenda for the conversation we’re already set to have. PublicLogic comes ready to build; "
     "this page is just so everyone’s working from the same picture going in.",before=4,after=3)
para("PublicLogic LLC · Prepared for discussion with WasteWerx. This outlines a proposed partnership; it is "
     "not a commitment or a guarantee of any award. Roles and budget are finalized in the application.  ·  "
     "nate@publiclogic.org",SANS,8,MUTE,ital=True,align=WD_ALIGN_PARAGRAPH.CENTER,before=2,after=0)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"final_deck","PublicLogic - WasteWerx Funding Path (STTR one-pager).docx")
os.makedirs(os.path.dirname(out),exist_ok=True)
d.save(out); print("wrote",out)
