#!/usr/bin/env python3
"""
PublicLogic — Permit Path Scan: 30-Day Cash Offer Run Book.
Tactical run book to generate cash this month through the public-facing
LogicCommons / Permit & Bridge layer (brand ring-fence keeps PublicLogic
stewardship uncheapened). Faithful to the uploaded source. Editable .docx.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SLATE=RGBColor(0x1A,0x1D,0x20); GREEN=RGBColor(0x0F,0x4C,0x3A)
GOLD=RGBColor(0xA9,0x77,0x2F); MUTE=RGBColor(0x6B,0x66,0x60); RED=RGBColor(0x8B,0x2A,0x1A)
SERIF="Georgia"; SANS="Calibri"

d=Document()
for s in d.sections:
    s.top_margin=Inches(0.6); s.bottom_margin=Inches(0.55)
    s.left_margin=Inches(0.85); s.right_margin=Inches(0.85)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hexc); tcPr.append(sh)

def para(text,font=SANS,size=10.5,color=SLATE,bold=False,ital=False,
         align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.06):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    r=p.add_run(text); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=color; f.bold=bold; f.italic=ital
    return p

def runs(p_runs,size=10.5,align=WD_ALIGN_PARAGRAPH.LEFT,before=2,after=2,sp=1.06):
    p=d.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=sp
    for (t,font,col,bold,ital) in p_runs:
        r=p.add_run(t); f=r.font; f.name=font; f.size=Pt(size); f.color.rgb=col; f.bold=bold; f.italic=ital
    return p

def h(text): para(text,SANS,11,GREEN,bold=True,before=9,after=2)
def bullet(text,g=GOLD):
    runs([("—  ",SANS,g,True,False),(text,SANS,SLATE,False,False)],after=1)

def table(headers,rows,widths,headfill="1A1D20"):
    t=d.add_table(rows=0,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    if any(headers):
        hr=t.add_row().cells
        for i,htxt in enumerate(headers):
            hr[i].width=Inches(widths[i]); shade(hr[i],headfill)
            r=hr[i].paragraphs[0].add_run(htxt); r.font.name=SANS; r.font.size=Pt(9.5); r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].width=Inches(widths[i])
            r=cells[i].paragraphs[0].add_run(val); r.font.name=SANS; r.font.size=Pt(10); r.font.color.rgb=SLATE
            if i==0: r.font.bold=True
    return t

# ---- Header
para("PUBLICLOGIC   ·   LOGICCOMMONS / PERMIT & BRIDGE",SANS,10,GOLD,bold=True,after=0)
para("Permit Path Scan — 30-Day Cash Offer",SERIF,23,SLATE,bold=True,before=0,after=0)
para("Run book · v1 · June 2026 · Private & Confidential.  Sells under the public-facing layer so "
     "PublicLogic stewardship stays uncheapened.",SANS,10,GREEN,ital=True,before=1,after=4)

# the one line
runs([("The one line:   ",SANS,GREEN,True,False),
      ("“Stop guessing before you spend money.”",SERIF,SLATE,True,True)],size=13,after=1)
para("Don’t sell stewardship for this. Sell relief from permitting confusion to people with a deadline "
     "and a checkbook.",SANS,9.5,MUTE,ital=True,after=3)

# ---- The Offer
h("THE OFFER")
runs([("Permit Path Scan — $500 flat, 3 business days. ",SANS,SLATE,True,False),
      ("For anyone stuck asking: “Can I do this, who do I call, and what do I need?”",SANS,SLATE,False,False)],after=1)
runs([("Deliverable: ",SANS,GREEN,True,False),
      ("a 2–4 page memo — likely boards/offices, the permit path, documents needed, risk points, a first "
       "call/email script, and the next three moves.",SANS,SLATE,False,False)],after=1)
runs([("Language is load-bearing. ",SANS,RED,True,False),
      ("Everything stays in likely / probably. The Scan points at the most likely path; the permitting "
       "authority always makes the final call. Don’t let a phone call talk you out of the hedge — it’s "
       "what keeps a $500 opinion from becoming a liability.",SANS,SLATE,False,True)],after=2)

# ---- Price ladder
h("PRICE LADDER — 30 DAYS ONLY")
table(["Tier","Price","What it is"],
 [["Basic Scan","$500","Standard turnaround, 3 business days."],
  ["Rush Scan","$750","48-hour turnaround."],
  ["Complex Scan (E&O gate)","$1,250","Multiple boards / properties. Highest dollars and reliance — do not sell until E&O is confirmed."],
  ["Follow-up call","$150","Walk the memo, answer questions, point to the next move."]],
 [2.0,0.9,4.0])
runs([("Before the $1,250 tier — this week, not after scan #7:  ",SANS,RED,True,False),
      ("Permit-expediting is a legitimate, legal service in MA, so the basic scan is low-risk. The "
       "exposure concentrates in the complex/developer tier. Confirm E&O (professional liability) "
       "coverage before taking a developer’s $1,250. Keep every scan in likely/probably; never cross "
       "into “yes, this use is allowed.” That’s the line between triage and a legal opinion.",
       SANS,SLATE,False,True)],size=9.5,before=2,after=2)

# ---- Target buyers
h("TARGET BUYERS — PEOPLE WITH ACTIVE PAIN")
para("Go after whoever has a deadline, a purchase, a lease, a contractor, or an opening date attached:",after=1)
for b in ["Contractors, architects, realtors",
          "Landlords and small business owners (tenant fit-outs, signage, change of use)",
          "Churches / nonprofits with property",
          "Homeowners doing ADUs, sheds, additions, food businesses, events",
          "Developers with small projects"]:
    bullet(b)

# ---- Where the list comes from
h("WHERE THE LIST COMES FROM — MINE YOUR NETWORK, DON’T COLD-BLAST")
runs([("Cold sends to trades you have no relationship with convert at ~1–3%. ",SANS,SLATE,True,False),
      ("The fast path is warm and referral-fed. Keep the municipal network ON — as the referral engine, "
       "not the buyer. Building inspectors, clerks, and the realtors and contractors you already know are "
       "the channel to private applicants who can pay this month. The ask is simple: ",SANS,SLATE,False,False),
      ("“Who do you know who’s stuck?”",SANS,GREEN,True,True)],after=2)

# ---- Outreach email
h("OUTREACH EMAIL — COPY · PASTE · PERSONALIZE THE [BRACKETS]")
para("Lead with the credential nothing else here uses: you were the Town Administrator who sat on the "
     "other side of that counter. That sells a $500 scan faster than “fixed-fee service.”",
     SANS,9.5,MUTE,ital=True,after=2)
for ln,bold in [
 ("Subject:  Quick permit-path help for stuck projects",True),
 ("Hi [Name],",False),
 ("I spent years as a Massachusetts Town Administrator — I’ve sat on the board side of the permitting "
  "counter and know how these offices actually decide. I’m now offering a small fixed-fee service for "
  "people stuck figuring out what permit path they’re really on.",False),
 ("For $500 I review the project and give a short Permit Path Scan: who to call first, what board or "
  "office likely handles it, the documents they’ll probably need, where it tends to stall, and the next "
  "three moves. Three business days.",False),
 ("Useful for additions, ADUs, small commercial uses, food businesses, events, churches/nonprofits, "
  "tenant fit-outs, and odd zoning questions.",False),
 ("Do you know anyone right now who’s stuck, delayed, or unsure where to start? Happy to take a quick "
  "look and tell them if it’s even a fit.",False),
 ("Best,",False),
 ("Nathan Boudreau, MPA / MCPPO", False),
 ("PublicLogic · nate@publiclogic.org · 978-807-0829",False),
]:
    para(ln,SANS,10,SLATE,bold=bold,after=1,sp=1.05)

# ---- Call script
h("CALL SCRIPT")
para("Open — set the frame",SANS,10.5,GOLD,bold=True,before=2,after=1)
runs([("“This isn’t legal advice and it doesn’t replace the town. I help you stop guessing, understand "
       "the likely path, and show up prepared. If it’s simple, I’ll tell you. If it’s messy, I’ll show "
       "you where it’s messy.”",SERIF,SLATE,False,True)],after=2)
para("Qualify — five questions",SANS,10.5,GOLD,bold=True,after=1)
for q in ["What are you trying to do?","Where is the property?",
          "Have you already talked to anyone — and what did they say?",
          "Is there a deadline, purchase, lease, contractor, or opening date attached?",
          "How many boards or properties does this likely touch? (basic vs. complex)"]:
    bullet(q)
para("Close",SANS,10.5,GOLD,bold=True,before=2,after=1)
runs([("“This is a good fit for a Permit Path Scan. It’s $500, fixed fee. I can turn it around in three "
       "business days once I have the address, a project description, and any documents you already have.”",
       SERIF,SLATE,False,True)],after=1)
runs([("If it’s multi-board or multi-property: ",SANS,RED,True,False),
      ("quote the $1,250 Complex Scan — only if E&O is in place.",SANS,SLATE,False,False)],size=9.5,after=2)

# ---- 30-day goal
h("30-DAY GOAL")
runs([("The fastest cash path isn’t convincing towns to buy continuity. ",SANS,SLATE,True,False),
      ("It’s helping people with real projects stop bleeding time in permitting confusion.",SANS,SLATE,False,False)],after=1)
table(["Scenario","Cash"],
 [["6 scans × $500","$3,000"],["10 scans × $500","$5,000"],["8 scans + 2 complex","$6,500+"]],
 [4.0,2.0])

# ---- pause / lead with
h("FOR 30 DAYS — PAUSE vs. LEAD WITH")
table(["Pause as the lead","Lead with instead"],
 [["Stewardship Map as the opening offer","$500 Permit Path Scan"],
  ["Municipal cold outreach as the main path","Network as referral engine → private buyers"],
  ["Abstract “institutional stewardship” language","“Stop guessing before you spend money.”"],
  ["Long decks","One email, one call, one memo"]],
 [3.0,3.0])

# ---- ring-fence
h("THE BRAND RING-FENCE")
runs([("You don’t pause PublicLogic — you route the scan through LogicCommons / Permit & Bridge, the layer "
       "built for public-facing transactional work. Same cash, no brand cost. ",SANS,SLATE,False,False),
      ("Stewardship stays intact for the municipal sale later; the Scan is the wedge and the proof, not "
       "the ceiling.",SANS,GREEN,True,True)],after=4)

para("PublicLogic LLC  ·  Nathan Boudreau, MPA / MCPPO  ·  LogicCommons / Permit & Bridge  ·  Private & Confidential",
     SANS,8.5,MUTE,align=WD_ALIGN_PARAGRAPH.CENTER,before=4,after=0)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"final_deck","PublicLogic - Permit Path Scan - 30-Day Cash Offer (Run Book).docx")
os.makedirs(os.path.dirname(out),exist_ok=True)
d.save(out); print("wrote",out)
